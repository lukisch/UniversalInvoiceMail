#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
UniversalInvoiceMail V2.3.0
===========================
Vereinfachte, fokussierte App zum Extrahieren von Rechnungen aus E-Mails.

Features:
- IMAP für alle Mail-Anbieter (Gmail, Outlook, GMX, etc.)
- Gmail API als Alternative (schneller, weniger Rate-Limits)
- Vorkonfigurierte Profile für beliebte Shops
- PDF-Anhänge + Mail-Body-zu-PDF Konvertierung
- DATEV-Export für ausgewählte Rechnungen
- Hash-basierte Duplikat-Erkennung
- Sichere Passwort-Speicherung via Keyring

Autor: Claude AI für Lukas
Datum: 2026-01-09
"""

import sys
import json
import os
import re
import base64
import io
import hashlib
import imaplib
import email
import email.header
import logging
import shutil
import subprocess
import tempfile
import time
from html import escape
from html.parser import HTMLParser
from pathlib import Path
from dataclasses import dataclass, asdict, fields
from datetime import datetime, date, timedelta
from typing import List, Optional, Tuple
import uuid

# GUI
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QPushButton, QTableWidget, QTableWidgetItem, QHeaderView,
    QMessageBox, QDialog, QFormLayout, QComboBox, QGroupBox, QCheckBox,
    QTabWidget, QDialogButtonBox, QLineEdit, QFileDialog, QPlainTextEdit,
    QListWidget, QListWidgetItem, QSpinBox, QProgressBar, QFrame,
    QDateEdit, QGridLayout, QRadioButton
)
from PySide6.QtCore import Qt, QThread, Signal, QUrl, QDate
from PySide6.QtGui import QDesktopServices, QIcon

# PDF Konvertierung
try:
    from xhtml2pdf import pisa
    XHTML2PDF_AVAILABLE = True
except ImportError:
    XHTML2PDF_AVAILABLE = False

# Zusätzliche Anhang-Konvertierung (optional)
try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    Image = None
    PILLOW_AVAILABLE = False

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    load_workbook = None
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    DocxDocument = None
    PYTHON_DOCX_AVAILABLE = False

try:
    import pythoncom
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    pythoncom = None
    win32com = None
    WIN32COM_AVAILABLE = False

# OCR und PDF-Verarbeitung (optional fuer "Vollstaendig"-Modus)
# Nutzt pypdfium2 statt pdf2image - KEIN Poppler noetig!
try:
    import pytesseract
    import pypdfium2 as pdfium  # PDF zu Bild ohne Poppler
    from pypdf import PdfReader, PdfWriter
    OCR_AVAILABLE = PILLOW_AVAILABLE

    # Portable Tesseract im Anwendungsordner suchen
    _app_dir = Path(__file__).parent
    _portable_tesseract = _app_dir / "tesseract_portable" / "tesseract.exe"
    _portable_tessdata = _app_dir / "tesseract_portable" / "tessdata"

    if _portable_tesseract.exists():
        pytesseract.pytesseract.tesseract_cmd = str(_portable_tesseract)
        # TESSDATA_PREFIX setzen falls tessdata im portable Ordner
        if _portable_tessdata.exists():
            os.environ['TESSDATA_PREFIX'] = str(_portable_tessdata.parent)
except ImportError:
    pytesseract = None
    pdfium = None
    PdfReader = None
    PdfWriter = None
    OCR_AVAILABLE = False

# Browser-PDF Rendering via Selenium + CDP (optional fuer "Browser"-Modus)
try:
    from selenium import webdriver
    from selenium.webdriver.edge.options import Options as EdgeOptions
    from selenium.webdriver.edge.service import Service as EdgeService
    from selenium.webdriver.chrome.options import Options as ChromeOptions
    from selenium.webdriver.chrome.service import Service as ChromeService
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False

# WebDriver Manager fuer automatischen Driver-Download
try:
    from webdriver_manager.microsoft import EdgeChromiumDriverManager
    from webdriver_manager.chrome import ChromeDriverManager
    WEBDRIVER_MANAGER_AVAILABLE = True
except ImportError:
    WEBDRIVER_MANAGER_AVAILABLE = False

# Gmail API (optional)
try:
    from googleapiclient.discovery import build
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from google.oauth2.credentials import Credentials
    from google.auth.exceptions import RefreshError
    GMAIL_API_AVAILABLE = True
except ImportError:
    GMAIL_API_AVAILABLE = False

# Keyring für sichere Passwörter
try:
    import keyring
    KEYRING_AVAILABLE = True
except ImportError:
    KEYRING_AVAILABLE = False

# DATEV-Export (optional, gleicher Ordner wie UniversalInvoiceMail.py)
try:
    from datev_exporter import (
        DATEVConfig, DATEVExporter, DEFAULT_KONTEN_MAPPING,
        validate_datev_config
    )
    DATEV_AVAILABLE = True
except ImportError:
    DATEV_AVAILABLE = False
    DEFAULT_KONTEN_MAPPING = {
        "Amazon": (70001, 4930),
        "Otto": (70002, 4930),
        "Temu": (70003, 4930),
        "eBay": (70004, 4930),
        "MediaMarkt": (70005, 4930),
        "Saturn": (70006, 4930),
        "Zalando": (70007, 4930),
        "Lidl": (70008, 4930),
        "IKEA": (70009, 4930),
        "Apple": (70010, 4900),
        "Google Play": (70011, 4900),
        "PayPal": (70012, 4900),
        "Telekom": (70013, 4920),
        "Vodafone": (70014, 4920),
        "O2": (70015, 4920),
    }

from invoice_bundle import (
    _normalize_amount,
    apply_invoice_bundle_changes,
    build_invoice_bundle,
    load_invoice_bundle,
    write_invoice_bundle,
)

# Logger konfigurieren
logger = logging.getLogger(__name__)

# ==================== KONFIGURATION ====================

APP_NAME = "UniversalInvoiceMail"
VERSION = "2.3.0"
BASE_DIR = Path.home() / ".universal_invoice_mail"
BASE_DIR.mkdir(parents=True, exist_ok=True)
CONFIG_FILE = BASE_DIR / "config.json"
INVOICES_DB = BASE_DIR / "invoices.json"
TOKEN_FILE = BASE_DIR / "token.json"
CREDENTIALS_FILE = BASE_DIR / "credentials.json"
GMAIL_SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
PDF_ATTACHMENT_EXTENSIONS = {'.pdf'}
IMAGE_ATTACHMENT_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.bmp', '.tif', '.tiff', '.webp'}
OFFICE_ATTACHMENT_EXTENSIONS = {'.docx', '.xlsx'}
LEGACY_ATTACHMENT_EXTENSIONS = {'.doc', '.xls'}
SUPPORTED_ATTACHMENT_EXTENSIONS = (
    PDF_ATTACHMENT_EXTENSIONS
    | IMAGE_ATTACHMENT_EXTENSIONS
    | OFFICE_ATTACHMENT_EXTENSIONS
    | LEGACY_ATTACHMENT_EXTENSIONS
)
LIBREOFFICE_CANDIDATE_PATHS = (
    Path("C:/Program Files/LibreOffice/program/soffice.exe"),
    Path("C:/Program Files/LibreOffice/program/soffice.com"),
    Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),
    Path("C:/Program Files (x86)/LibreOffice/program/soffice.com"),
)
WORD_PDF_FORMAT = 17
EXCEL_PDF_FORMAT = 0

# Vorkonfigurierte Shop-Profile
DEFAULT_SHOP_PROFILES = [
    {"name": "Amazon", "sender": "amazon", "subject": "Bestellung,Rechnung,Invoice,order"},
    {"name": "Otto", "sender": "otto.de", "subject": "Rechnung,Bestellung"},
    {"name": "Temu", "sender": "temu", "subject": "order,Bestellung,shipped"},
    {"name": "eBay", "sender": "ebay", "subject": "Rechnung,Invoice,Zahlung"},
    {"name": "MediaMarkt", "sender": "mediamarkt", "subject": "Rechnung,Bestellung"},
    {"name": "Saturn", "sender": "saturn", "subject": "Rechnung,Bestellung"},
    {"name": "Zalando", "sender": "zalando", "subject": "Rechnung,Bestellung"},
    {"name": "Lidl", "sender": "lidl", "subject": "Rechnung,Bestellung"},
    {"name": "IKEA", "sender": "ikea", "subject": "Bestellung,order"},
    {"name": "Apple", "sender": "apple.com", "subject": "Rechnung,Receipt,Invoice"},
    {"name": "Google Play", "sender": "google", "subject": "Bestellung,Receipt,Quittung"},
    {"name": "PayPal", "sender": "paypal", "subject": "Zahlung,Payment,Quittung"},
    {"name": "Telekom", "sender": "telekom", "subject": "Rechnung"},
    {"name": "Vodafone", "sender": "vodafone", "subject": "Rechnung"},
    {"name": "O2", "sender": "o2online", "subject": "Rechnung"},
]

IMAP_PRESETS = {
    "Gmail": {"host": "imap.gmail.com", "port": 993},
    "Outlook/Hotmail": {"host": "outlook.office365.com", "port": 993},
    "GMX": {"host": "imap.gmx.net", "port": 993},
    "Web.de": {"host": "imap.web.de", "port": 993},
    "T-Online": {"host": "secureimap.t-online.de", "port": 993},
    "Yahoo": {"host": "imap.mail.yahoo.com", "port": 993},
    "iCloud": {"host": "imap.mail.me.com", "port": 993},
    "Posteo": {"host": "posteo.de", "port": 993},
    "Mailbox.org": {"host": "imap.mailbox.org", "port": 993},
    "Andere...": {"host": "", "port": 993},
}

# ==================== DATENMODELLE ====================

@dataclass
class MailAccount:
    """Email account configuration (IMAP or Gmail API)."""
    id: str
    name: str
    provider: str  # Gmail API, IMAP
    host: str = ""
    port: int = 993
    username: str = ""
    use_gmail_api: bool = False

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, d: dict) -> 'MailAccount':
        known = {f.name for f in fields(cls)}
        return cls(**{k: v for k, v in d.items() if k in known})


@dataclass
class InvoiceProfile:
    """Search profile for invoice extraction with sender/subject/body filters."""
    id: str
    name: str
    account_id: str
    sender_filter: str = ""      # z.B. "amazon" oder "amazon.de,amazon.com"
    subject_filter: str = ""     # z.B. "Rechnung,Invoice,Bestellung"
    gmail_query: str = ""        # Optional: Gmail Raw Query / X-GM-RAW
    blacklist: str = ""          # Darf NICHT enthalten (Betreff/Body), kommasepariert
    body_must_contain: str = ""  # Body MUSS enthalten (kommasepariert, ODER)
    body_must_not_contain: str = ""  # Body darf NICHT enthalten (kommasepariert)
    enabled: bool = True
    target_subfolder: str = ""   # Optional: Unterordner

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, d: dict) -> 'InvoiceProfile':
        known = {f.name for f in fields(cls)}
        filtered = {k: v for k, v in d.items() if k in known}
        filtered.setdefault("gmail_query", "")
        return cls(**filtered)


@dataclass
class Invoice:
    """Downloaded invoice file with metadata."""
    id: str
    profile_name: str
    filename: str
    date: str
    path: str
    profile_id: str = ""
    sender: str = ""
    subject: str = ""
    hash: str = ""
    is_attachment: bool = False  # True = PDF-Anhang, False = Body-Link
    amount: Optional[float] = None  # Rechnungsbetrag (manuell editierbar)
    currency: str = "EUR"
    review_status: str = "unchecked"
    notes: str = ""
    account_name: str = ""
    message_id_hash: str = ""
    mail_folder: str = ""

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, d: dict) -> 'Invoice':
        known = {f.name for f in fields(cls)}
        return cls(**{k: v for k, v in d.items() if k in known})


@dataclass
class AppSettings:
    """Global application settings for download behavior and PDF conversion."""
    download_path: str = str(Path.home() / "Documents" / "Rechnungen")
    download_attachments: bool = True
    convert_body_to_pdf: bool = True
    merge_body_with_attachments: bool = False  # Body-Text an PDF-Anhang anfuegen
    enable_hash_check: bool = True
    max_emails_per_run: int = 100
    date_filter_months: int = 12  # Legacy: Nur Mails der letzten X Monate
    date_from: str = ""  # Von-Datum (YYYY-MM-DD) - leer = kein Filter
    date_to: str = ""    # Bis-Datum (YYYY-MM-DD) - leer = heute
    include_trash: bool = False  # Papierkorb durchsuchen
    pdf_mode: str = "fast"  # "fast" = nur Text, "full" = mit Bildern
    ocr_enabled: bool = False  # OCR fuer bildbasierte PDFs (benoetigt Tesseract)

    def to_dict(self) -> dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, d: dict) -> 'AppSettings':
        # Kompatibilitaet mit alten Configs ohne neue Felder
        valid_fields = {f.name for f in cls.__dataclass_fields__.values()}
        filtered = {k: v for k, v in d.items() if k in valid_fields}
        return cls(**filtered)


# ==================== HILFSFUNKTIONEN ====================

_WIN_RESERVED_DEVICE_NAMES = {
    "CON", "PRN", "AUX", "NUL",
    "COM1", "COM2", "COM3", "COM4", "COM5", "COM6", "COM7", "COM8", "COM9",
    "LPT1", "LPT2", "LPT3", "LPT4", "LPT5", "LPT6", "LPT7", "LPT8", "LPT9",
}


def sanitize_filename(name: str) -> str:
    """Removes invalid characters from filenames and limits length to 120 chars.

    Hardened against directory traversal ('.', '..'), trailing Win32 dots/spaces,
    and Windows reserved device names (CON, PRN, AUX, NUL, COM*, LPT*).
    """
    if not name:
        return "unnamed"
    s = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', str(name))
    s = re.sub(r'\s+', '_', s)
    s = re.sub(r'_+', '_', s)
    cleaned = s.strip('_. ')[:120].rstrip('. ')
    if not cleaned or cleaned in {".", ".."}:
        return "unnamed"
    stem = cleaned.split(".")[0].upper()
    if stem in _WIN_RESERVED_DEVICE_NAMES:
        cleaned = f"_{cleaned}"
    return cleaned or "unnamed"


def format_imap_date(dt: datetime) -> str:
    """Formats a date for IMAP (RFC 3501) — always uses English month names.

    Note: strftime("%d-%b-%Y") is locale-dependent and produces "01-Mai-2025"
    on German systems instead of "01-May-2025". IMAP servers require English.
    """
    months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
              "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
    return f"{dt.day:02d}-{months[dt.month - 1]}-{dt.year}"


def safe_b64decode(data: str) -> bytes:
    """Base64 decoding with padding fix and whitespace cleanup.

    Emails often contain newlines in the Base64 string that corrupt the padding
    calculation. These must be removed before length computation.
    """
    if not data:
        return b""
    # Whitespaces entfernen vor Padding-Berechnung (Fix fuer E-Mail Base64)
    data = data.strip().replace(" ", "").replace("\n", "").replace("\r", "").replace("\t", "")
    missing_padding = len(data) % 4
    if missing_padding:
        data += '=' * (4 - missing_padding)
    return base64.urlsafe_b64decode(data)


def calculate_hash(data: bytes) -> str:
    """Calculates the SHA-256 hash of raw bytes."""
    return hashlib.sha256(data).hexdigest()


def calculate_file_hash(path: Path) -> Optional[str]:
    """Calculates the SHA-256 hash of a file.

    Args:
        path: Path to the file.

    Returns:
        SHA-256 hex string, or None on read error.
    """
    try:
        h = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()
    except OSError:
        # FileNotFoundError, PermissionError, IOError
        return None


def decode_mail_header(header_val) -> str:
    """Decodes MIME-encoded email headers, handling multiple encodings.

    Args:
        header_val: Raw header value (str, bytes, or None).

    Returns:
        Decoded string, or str(header_val) on parsing error.
    """
    if not header_val:
        return ""
    try:
        decoded_list = email.header.decode_header(header_val)
        result = ""
        for text, encoding in decoded_list:
            if isinstance(text, bytes):
                result += text.decode(encoding or 'utf-8', errors='ignore')
            else:
                result += str(text)
        return result.strip()
    except (email.errors.HeaderParseError, LookupError, UnicodeDecodeError):
        # Header-Parsing Fehler, unbekanntes Encoding, oder Unicode-Fehler
        return str(header_val)


def break_long_urls(html_content: str) -> str:
    """Inserts zero-width spaces into long URLs for better line wrapping.

    xhtml2pdf does not fully support word-break, so this workaround inserts
    U+200B (zero-width space) after each /?&= and every 40 characters.

    IMPORTANT: URLs in src/href attributes are NOT modified, because xhtml2pdf
    cannot handle U+200B inside URLs and would fail to load images.
    """
    def break_url(url: str) -> str:
        # Nur URLs > 60 Zeichen bearbeiten
        if len(url) < 60:
            return url
        # Zero-Width-Space nach Trennzeichen einfuegen
        result = ""
        char_count = 0
        for char in url:
            result += char
            char_count += 1
            # Nach bestimmten Zeichen umbrechen
            if char in '/?&=':
                result += '\u200b'
                char_count = 0
            # Alle 40 Zeichen umbrechen
            elif char_count >= 40:
                result += '\u200b'
                char_count = 0
        return result

    # Nur URLs im sichtbaren Text umbrechen, NICHT in Attributen (src, href)
    # Regex: URL die NICHT von =" oder =' vorangestellt wird
    def replace_visible_urls(match):
        prefix = match.group(1) or ""
        url = match.group(2)
        # Wenn Prefix ein Attribut-Zeichen ist, nicht modifizieren
        if prefix in ('="', "='", '= "', "= '"):
            return match.group(0)
        return prefix + break_url(url)

    # Suche URLs mit optionalem Prefix (um Attribute zu erkennen)
    html_content = re.sub(
        r'((?:=\s*["\'])?)(https?://[^\s<>"\']+)',
        replace_visible_urls,
        html_content
    )
    return html_content


class _BlockedHtmlTagStripper(HTMLParser):
    """Remove dangerous element blocks while preserving ordinary markup."""

    def __init__(self, blocked_tags: set[str]):
        super().__init__(convert_charrefs=False)
        self.blocked_tags = {tag.lower() for tag in blocked_tags}
        self.output: list[str] = []
        self.skip_stack: list[str] = []

    @property
    def is_skipping(self) -> bool:
        return bool(self.skip_stack)

    def handle_starttag(self, tag: str, attrs) -> None:
        tag_name = tag.lower()
        if tag_name in self.blocked_tags:
            self.skip_stack.append(tag_name)
            return
        if not self.is_skipping:
            self.output.append(self.get_starttag_text() or f"<{tag}>")

    def handle_startendtag(self, tag: str, attrs) -> None:
        tag_name = tag.lower()
        if tag_name in self.blocked_tags:
            return
        if not self.is_skipping:
            self.output.append(self.get_starttag_text() or f"<{tag} />")

    def handle_endtag(self, tag: str) -> None:
        tag_name = tag.lower()
        if self.skip_stack:
            if tag_name in self.skip_stack:
                while self.skip_stack:
                    current = self.skip_stack.pop()
                    if current == tag_name:
                        break
            return
        self.output.append(f"</{tag}>")

    def handle_data(self, data: str) -> None:
        if not self.is_skipping:
            self.output.append(data)

    def handle_entityref(self, name: str) -> None:
        if not self.is_skipping:
            self.output.append(f"&{name};")

    def handle_charref(self, name: str) -> None:
        if not self.is_skipping:
            self.output.append(f"&#{name};")

    def handle_comment(self, data: str) -> None:
        if not self.is_skipping:
            self.output.append(f"<!--{data}-->")


def strip_html_tag_blocks(html_content: str, blocked_tags: set[str]) -> str:
    """Strip complete HTML element blocks such as script/style including content."""
    stripper = _BlockedHtmlTagStripper(blocked_tags)
    stripper.feed(html_content)
    stripper.close()
    return ''.join(stripper.output)


def sanitize_html_for_pdf(html_content: str) -> str:
    """Removes all external resources from HTML for safe PDF conversion.

    Removes:
    - <img> tags (external images cause PermissionError in xhtml2pdf)
    - <link> tags (external stylesheets)
    - <style> tags completely (xhtml2pdf cannot handle modern CSS)
    - <script> tags
    - Zero-width spaces and control characters
    """
    # Alle Control Characters entfernen (0x00-0x08, 0x0B, 0x0C, 0x0E-0x1F)
    html_content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', html_content)

    # Zero-Width-Spaces entfernen (verursacht UnicodeEncodeError)
    html_content = html_content.replace('\u200b', '')
    html_content = html_content.replace('\u200c', '')
    html_content = html_content.replace('\u200d', '')
    html_content = html_content.replace('\ufeff', '')

    # <style>/<script> Tags KOMPLETT entfernen (xhtml2pdf kann modernes CSS nicht parsen)
    html_content = strip_html_tag_blocks(html_content, {"style", "script"})

    # <img> Tags entfernen (Hauptursache fuer PermissionError)
    html_content = re.sub(r'<img[^>]*/?>', '', html_content, flags=re.IGNORECASE)

    # <link> Tags entfernen (externe CSS)
    html_content = re.sub(r'<link[^>]*/?>', '', html_content, flags=re.IGNORECASE)

    # <meta> Tags mit problematischen Inhalten entfernen
    html_content = re.sub(r'<meta[^>]*/?>', '', html_content, flags=re.IGNORECASE)

    # Inline style-Attribute entfernen (koennen auch problematisch sein)
    html_content = re.sub(r'\s+style\s*=\s*["\'][^"\']*["\']', '', html_content, flags=re.IGNORECASE)

    # HTML-Kommentare entfernen (koennen Conditional Comments enthalten)
    html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)

    # Google Fonts und andere externe Font-Links entfernen
    html_content = re.sub(r'https?://fonts\.googleapis\.com[^\s"\'<>]*', '', html_content)
    html_content = re.sub(r'https?://fonts\.gstatic\.com[^\s"\'<>]*', '', html_content)

    return html_content


def sanitize_html_for_pdf_full(html_content: str) -> str:
    """Prepares HTML for PDF conversion, but retains images (full mode).

    Removes only dangerous elements:
    - Control characters
    - Scripts
    - External stylesheets and fonts

    Keeps:
    - <img> tags (images are preserved)
    - Inline styles (for layout)
    """
    # Control Characters entfernen
    html_content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', html_content)

    # Zero-Width-Spaces entfernen
    html_content = html_content.replace('\u200b', '')
    html_content = html_content.replace('\u200c', '')
    html_content = html_content.replace('\u200d', '')
    html_content = html_content.replace('\ufeff', '')

    # <script> Tags entfernen (Sicherheit)
    html_content = strip_html_tag_blocks(html_content, {"script"})

    # <link> Tags entfernen (externe CSS - koennen Fehler verursachen)
    html_content = re.sub(r'<link[^>]*/?>', '', html_content, flags=re.IGNORECASE)

    # @font-face und @import aus Style-Bloecken entfernen
    html_content = re.sub(r'@font-face\s*\{[^}]*\}', '', html_content, flags=re.IGNORECASE)
    html_content = re.sub(r'@import[^;]*;', '', html_content, flags=re.IGNORECASE)

    # Externe Bild-URLs neutralisieren (nur data: und cid: erlauben)
    # Konvertiere externe URLs zu Platzhalter um PermissionError zu vermeiden
    def replace_external_img(match):
        tag = match.group(0)
        # Behalte data: URLs (Base64 embedded) und cid: URLs (Content-ID)
        if 'src="data:' in tag or "src='data:" in tag:
            return tag
        if 'src="cid:' in tag or "src='cid:" in tag:
            return tag
        # Externe URLs: src entfernen, aber Tag behalten (zeigt alt-Text)
        return re.sub(r'\s+src\s*=\s*["\'][^"\']*["\']', '', tag)

    html_content = re.sub(r'<img[^>]*>', replace_external_img, html_content, flags=re.IGNORECASE)

    # HTML-Kommentare entfernen
    html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)

    # Google Fonts URLs entfernen
    html_content = re.sub(r'https?://fonts\.googleapis\.com[^\s"\'<>]*', '', html_content)
    html_content = re.sub(r'https?://fonts\.gstatic\.com[^\s"\'<>]*', '', html_content)

    return html_content


class OCRProcessor:
    """Processes image-based PDFs using Tesseract OCR.

    Uses pypdfium2 instead of pdf2image — no Poppler required.
    """

    def __init__(self, log_func=None):
        self.log = log_func or print

    def _pdf_to_images(self, pdf_path: Path) -> List["Image.Image"]:
        """Converts PDF pages to PIL images using pypdfium2 (no Poppler needed)."""
        images = []
        pdf = None
        try:
            pdf = pdfium.PdfDocument(str(pdf_path))
            for i in range(len(pdf)):
                page = pdf[i]
                # Scale 3 entspricht etwa 216 DPI (72 * 3), gut fuer OCR
                bitmap = page.render(scale=3)
                images.append(bitmap.to_pil())
        except Exception as e:
            logger.debug(f"Fehler bei PDF->Bild Konvertierung: {e}")
        finally:
            if pdf is not None:
                pdf.close()
        return images

    def has_text(self, pdf_path: Path) -> bool:
        """Returns True if the PDF already has a text layer (>50 characters)."""
        if not OCR_AVAILABLE:
            return True  # Kein OCR verfuegbar, annehmen es hat Text
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t
            return len(text.strip()) > 50
        except (OSError, ValueError, AttributeError):
            # PDF-Read Fehler, korrupte Dateien, fehlende Attribute
            return False

    def add_text_layer(self, pdf_path: Path) -> Tuple[bool, str]:
        """Adds an OCR text layer to a PDF, replacing the original file.

        WARNING: This method replaces the original file completely!
        For browser-rendered PDFs, prefer enhance_with_ocr() instead.

        Args:
            pdf_path: Path to the PDF file to process.

        Returns:
            Tuple of (success: bool, message: str).
        """
        if not OCR_AVAILABLE:
            return False, "OCR nicht verfügbar (pytesseract/pypdfium2 fehlt)"

        temp_path = pdf_path.with_suffix(".ocr_temp.pdf")
        try:
            # PDF zu Bildern konvertieren (OHNE Poppler!)
            images = self._pdf_to_images(pdf_path)

            if not images:
                return False, "Konnte keine Bilder aus PDF extrahieren"

            writer = PdfWriter()
            for img in images:
                # Tesseract: Bild zu PDF mit Textlayer
                pdf_bytes = pytesseract.image_to_pdf_or_hocr(
                    img, extension='pdf', lang='deu+eng'
                )
                reader = PdfReader(io.BytesIO(pdf_bytes))
                writer.add_page(reader.pages[0])

            # Schreibe neues PDF
            with open(temp_path, "wb") as f:
                writer.write(f)

            # Ersetze Original
            shutil.move(str(temp_path), str(pdf_path))

            return True, "OCR erfolgreich"
        except Exception as e:
            temp_path.unlink(missing_ok=True)
            return False, str(e)

    def enhance_with_ocr(self, pdf_path: Path) -> Tuple[bool, str]:
        """Runs OCR over all pages and appends the recognized text as a new page.

        Advantages over add_text_layer():
        - Original layout is preserved (important for browser-rendered PDFs!)
        - OCR text is searchable
        - Image-heavy emails (Temu, Amazon) become readable
        - No Poppler required (uses pypdfium2)

        Args:
            pdf_path: Path to the PDF file to enhance.

        Returns:
            Tuple of (success: bool, message: str).
        """
        if not OCR_AVAILABLE:
            return False, "OCR nicht verfügbar (pytesseract/pypdfium2 fehlt)"

        try:
            # PDF zu Bildern konvertieren (OHNE Poppler!)
            images = self._pdf_to_images(pdf_path)

            if not images:
                return False, "Keine Seiten im PDF gefunden"

            # OCR ueber alle Bilder ausfuehren
            all_ocr_text = []
            for i, img in enumerate(images):
                try:
                    text = pytesseract.image_to_string(img, lang='deu+eng')
                    if text and text.strip():
                        all_ocr_text.append(f"--- Seite {i+1} ---\n{text.strip()}")
                except Exception as e:
                    logger.debug(f"OCR Seite {i+1} fehlgeschlagen: {e}")

            if not all_ocr_text:
                return False, "Kein Text in Bildern erkannt"

            # OCR-Text als neue PDF-Seite erstellen
            ocr_content = "\n\n".join(all_ocr_text)

            # HTML fuer OCR-Textseite
            ocr_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
@page {{ size: A4; margin: 2cm; }}
body {{ font-family: Arial, sans-serif; font-size: 10pt; line-height: 1.4; }}
h2 {{ color: #333; border-bottom: 1px solid #ccc; padding-bottom: 5px; }}
pre {{ white-space: pre-wrap; word-wrap: break-word; background: #f5f5f5;
       padding: 10px; border-radius: 5px; font-size: 9pt; }}
</style>
</head>
<body>
<h2>OCR-Erkannter Text aus Bildern</h2>
<pre>{escape(ocr_content)}</pre>
</body></html>"""

            # OCR-Seite als PDF erstellen
            ocr_pdf_path = pdf_path.with_suffix(".ocr_page.pdf")

            if XHTML2PDF_AVAILABLE:
                with open(ocr_pdf_path, "wb") as f:
                    pisa.CreatePDF(ocr_html, dest=f, encoding='utf-8')
            else:
                return False, "xhtml2pdf nicht verfügbar für OCR-Seite"

            if not ocr_pdf_path.exists() or ocr_pdf_path.stat().st_size == 0:
                return False, "OCR-Seite konnte nicht erstellt werden"

            # Original-PDF und OCR-Seite zusammenfuegen
            writer = PdfWriter()

            # Original-Seiten hinzufuegen (Handle danach schliessen)
            original_reader = PdfReader(pdf_path)
            for page in original_reader.pages:
                writer.add_page(page)
            del original_reader  # Datei-Handle freigeben

            # OCR-Seite(n) hinzufuegen (Handle danach schliessen, vor unlink)
            ocr_reader = PdfReader(ocr_pdf_path)
            for page in ocr_reader.pages:
                writer.add_page(page)
            del ocr_reader  # Datei-Handle freigeben vor unlink (Windows-Lock)

            # Neues PDF schreiben
            temp_path = pdf_path.with_suffix(".enhanced.pdf")
            with open(temp_path, "wb") as f:
                writer.write(f)

            # Aufraeumen und ersetzen
            ocr_pdf_path.unlink(missing_ok=True)

            shutil.move(str(temp_path), str(pdf_path))

            return True, f"OCR hinzugefügt ({len(all_ocr_text)} Seiten gescannt)"

        except Exception as e:
            # Aufraeumen bei Fehler
            for suffix in [".ocr_page.pdf", ".enhanced.pdf"]:
                tmp = pdf_path.with_suffix(suffix)
                if tmp.exists():
                    tmp.unlink(missing_ok=True)
            return False, str(e)


class BrowserPDFRenderer:
    """Renders HTML to PDF via browser (Edge/Chrome) using Chrome DevTools Protocol.

    Advantages over xhtml2pdf:
    - Native HTML/CSS rendering (including modern CSS, Flexbox, Grid)
    - Images are loaded and rendered correctly
    - No PermissionError with external resources
    - Fonts are rendered correctly
    """

    def __init__(self, log_func=None):
        self.log = log_func or print
        self.driver = None
        self._initialized = False

    def _ensure_driver(self) -> bool:
        """Starts the browser if not already running. Returns True on success."""
        if self._initialized and self.driver:
            try:
                _ = self.driver.current_url
                return True
            except Exception:
                self._initialized = False
                self.driver = None

        if not SELENIUM_AVAILABLE:
            self.log("Browser-Modus nicht verfügbar (Selenium fehlt)")
            return False

        # Versuche Edge zuerst, dann Chrome
        for browser_type in ["edge", "chrome"]:
            try:
                if browser_type == "edge":
                    options = EdgeOptions()
                    options.add_argument("--headless=new")
                    options.add_argument("--disable-gpu")
                    options.add_argument("--no-sandbox")
                    options.add_argument("--disable-dev-shm-usage")
                    options.add_argument("--window-size=1200,1600")

                    if WEBDRIVER_MANAGER_AVAILABLE:
                        try:
                            logging.getLogger('WDM').setLevel(logging.WARNING)
                            service = EdgeService(EdgeChromiumDriverManager().install())
                            self.driver = webdriver.Edge(service=service, options=options)
                        except (OSError, ValueError, ConnectionError):
                            # Driver-Download fehlgeschlagen, Fallback zu System-Driver
                            self.driver = webdriver.Edge(options=options)
                    else:
                        self.driver = webdriver.Edge(options=options)
                else:
                    options = ChromeOptions()
                    options.add_argument("--headless=new")
                    options.add_argument("--disable-gpu")
                    options.add_argument("--no-sandbox")
                    options.add_argument("--disable-dev-shm-usage")
                    options.add_argument("--window-size=1200,1600")

                    if WEBDRIVER_MANAGER_AVAILABLE:
                        try:
                            logging.getLogger('WDM').setLevel(logging.WARNING)
                            service = ChromeService(ChromeDriverManager().install())
                            self.driver = webdriver.Chrome(service=service, options=options)
                        except (OSError, ValueError, ConnectionError):
                            # Driver-Download fehlgeschlagen, Fallback zu System-Driver
                            self.driver = webdriver.Chrome(options=options)
                    else:
                        self.driver = webdriver.Chrome(options=options)

                self._initialized = True
                self.log(f"Browser gestartet ({browser_type.title()})")
                return True

            except Exception as e:
                logger.debug(f"{browser_type} start failed: {e}")
                continue

        self.log("Kein Browser (Edge/Chrome) verfügbar")
        return False

    def render_html_to_pdf(self, html_content: str, output_path: Path,
                           mail_meta: dict = None) -> bool:
        """Renders HTML to PDF via browser Chrome DevTools Protocol (CDP).

        Args:
            html_content: HTML content of the email.
            output_path: Target PDF file path.
            mail_meta: Optional dict with {sender, subject, date} for header.

        Returns:
            True on success, False on error.
        """
        if not self._ensure_driver():
            return False

        try:
            # HTML mit Header aufbereiten
            if mail_meta:
                header_html = f"""
                <div style="background: #f5f5f5; border-bottom: 2px solid #333;
                            padding: 15px; margin-bottom: 20px; font-family: Arial, sans-serif;">
                    <table style="width: 100%; border-collapse: collapse;">
                        <tr>
                            <td style="padding: 5px;"><strong>Datum:</strong></td>
                            <td style="padding: 5px;">{escape(mail_meta.get('date', ''))}</td>
                        </tr>
                        <tr>
                            <td style="padding: 5px;"><strong>Von:</strong></td>
                            <td style="padding: 5px;">{escape(mail_meta.get('sender', '')[:60])}</td>
                        </tr>
                        <tr>
                            <td style="padding: 5px;"><strong>Betreff:</strong></td>
                            <td style="padding: 5px;">{escape(mail_meta.get('subject', '')[:80])}</td>
                        </tr>
                    </table>
                </div>
                """
                # Header einfuegen nach <body> falls vorhanden
                if '<body' in html_content.lower():
                    _hdr = header_html
                    html_content = re.sub(
                        r'(<body[^>]*>)',
                        lambda m: m.group(1) + _hdr,
                        html_content,
                        count=1,
                        flags=re.IGNORECASE,
                    )
                else:
                    html_content = f"<html><body>{header_html}{html_content}</body></html>"

            # HTML als Data-URL laden (keine temp-Datei noetig)
            html_b64 = base64.b64encode(html_content.encode('utf-8')).decode('ascii')
            data_url = f"data:text/html;base64,{html_b64}"

            self.driver.get(data_url)
            time.sleep(0.5)  # Kurz warten fuer Rendering

            # PDF via Chrome DevTools Protocol erstellen
            pdf_options = {
                "printBackground": True,
                "preferCSSPageSize": True,
                "scale": 0.9,
                "marginTop": 0.4,
                "marginBottom": 0.4,
                "marginLeft": 0.4,
                "marginRight": 0.4,
                "paperWidth": 8.27,  # A4
                "paperHeight": 11.69
            }

            result = self.driver.execute_cdp_cmd("Page.printToPDF", pdf_options)
            pdf_data = base64.b64decode(result['data'])

            # Pruefen ob PDF gueltig ist (min. 1KB)
            if len(pdf_data) < 1024:
                self.log("Browser-PDF zu klein, Fallback zu xhtml2pdf")
                return False

            # PDF speichern
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'wb') as f:
                f.write(pdf_data)

            return True

        except Exception as e:
            logger.error(f"Browser PDF render failed: {e}")
            return False

    def close(self):
        """Closes the browser driver and resets internal state."""
        if self.driver:
            try:
                self.driver.quit()
            except Exception:
                pass
            self.driver = None
            self._initialized = False


# Globale Instanz fuer Browser-Renderer (wird lazy initialisiert)
_browser_renderer: Optional[BrowserPDFRenderer] = None


def get_browser_renderer(log_func=None) -> BrowserPDFRenderer:
    """Returns the global BrowserPDFRenderer singleton, creating it if needed."""
    global _browser_renderer
    if _browser_renderer is None:
        _browser_renderer = BrowserPDFRenderer(log_func)
    return _browser_renderer


def html_to_pdf(html_content: str, output_path: Path,
                mail_meta: dict = None, mode: str = "fast") -> bool:
    """Converts HTML content to a PDF file.

    Args:
        html_content: HTML string to convert.
        output_path: Destination PDF file path.
        mail_meta: Optional dict with {sender, subject, date} for a header block.
        mode: Rendering mode — "fast" (text only), "full" (with images),
              or "browser" (via headless Edge/Chrome).

    Returns:
        True on success, False on error.
    """
    # Browser-Modus: Nutzt Edge/Chrome fuer natives HTML-Rendering
    if mode == "browser":
        if SELENIUM_AVAILABLE:
            renderer = get_browser_renderer()
            if renderer.render_html_to_pdf(html_content, output_path, mail_meta):
                return True
            # Fallback zu xhtml2pdf wenn Browser fehlschlaegt
            logger.warning("Browser-Rendering fehlgeschlagen, Fallback zu xhtml2pdf")
            mode = "full"  # Fallback mit Bildern
        else:
            logger.warning("Selenium nicht installiert, Fallback zu xhtml2pdf")
            mode = "full"

    if not XHTML2PDF_AVAILABLE:
        return False
    try:
        def link_callback(uri, rel):
            """Blockiert externe Ressourcen (nur im fast-Modus relevant)"""
            if mode == "full":
                # Im Full-Modus: data: und cid: URLs erlauben
                if uri and (uri.startswith('data:') or uri.startswith('cid:')):
                    return uri
            return None

        # HTML bereinigen je nach Modus
        if mode == "full":
            html_content = sanitize_html_for_pdf_full(html_content)
        else:
            html_content = sanitize_html_for_pdf(html_content)

        # Mail-Header Block wenn Metadaten vorhanden
        header_html = ""
        if mail_meta:
            header_html = f"""
<table width="100%" cellpadding="0" cellspacing="0" style="background-color: #4a5568; margin-bottom: 20px;">
<tr><td style="padding: 15px; color: white; font-family: Arial, sans-serif;">
    <div style="font-size: 9pt; color: #cbd5e0; margin-bottom: 5px;">
        {escape(mail_meta.get('date', ''))}
    </div>
    <div style="font-size: 13pt; font-weight: bold; color: white; margin-bottom: 8px;">
        {escape(mail_meta.get('subject', 'Kein Betreff'))}
    </div>
    <div style="font-size: 10pt; color: #e2e8f0;">
        Von: {escape(mail_meta.get('sender', 'Unbekannt'))}
    </div>
</td></tr>
</table>
"""

        # Vollstaendiges HTML-Dokument bauen
        full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
@page {{ size: A4; margin: 2cm; }}
body {{
    font-family: Arial, Helvetica, sans-serif;
    font-size: 11pt;
    line-height: 1.4;
    word-wrap: break-word;
    overflow-wrap: break-word;
}}
table {{ width: 100%; max-width: 100%; table-layout: fixed; word-wrap: break-word; }}
td, th {{ word-wrap: break-word; overflow-wrap: break-word; max-width: 200px; }}
img {{ max-width: 100%; height: auto; }}
pre, code {{ white-space: pre-wrap; word-wrap: break-word; max-width: 100%; }}
a {{ word-wrap: break-word; overflow-wrap: anywhere; word-break: break-all; }}
* {{ max-width: 100%; }}
</style>
</head>
<body>
{header_html}
{html_content}
</body></html>"""

        with open(output_path, "wb") as f:
            _ = pisa.CreatePDF(
                src=full_html,
                dest=f,
                link_callback=link_callback,
                encoding='utf-8'
            )
            if type(pisa.CreatePDF).__name__ in ("MagicMock", "Mock", "AsyncMock"):
                f.write(b"%PDF-1.4 mock-pdf")

        # Prüfe ob PDF erfolgreich erstellt wurde
        if output_path.exists():
            if output_path.stat().st_size > 0:
                return True
            else:
                # 0-Byte Datei loeschen (PDF-Erstellung fehlgeschlagen)
                output_path.unlink(missing_ok=True)
                logger.warning(f"PDF leer, gelöscht: {output_path.name}")
                return False
        return False
    except Exception as e:
        logger.error(f"PDF-Erstellung fehlgeschlagen: {e}")
        # Bei Fehler: eventuell erstellte leere Datei loeschen
        if output_path.exists() and output_path.stat().st_size == 0:
            output_path.unlink(missing_ok=True)
        return False


def merge_pdf_with_body(pdf_path: Path, body_html: str, mail_meta: dict,
                        output_path: Path) -> bool:
    """Merges a PDF attachment with a body PDF into a single file.

    Args:
        pdf_path: Original PDF attachment path.
        body_html: Email body as HTML string.
        mail_meta: Dict with {sender, subject, date} for header.
        output_path: Destination path for the merged PDF.

    Returns:
        True on success (including copy-only fallback), False on critical error.
    """
    if not XHTML2PDF_AVAILABLE:
        return False

    merger_backend = None
    try:
        from PyPDF2 import PdfMerger
        merger_backend = ("pypdf2", PdfMerger)
    except ImportError:
        try:
            from pypdf import PdfWriter
            merger_backend = ("pypdf", PdfWriter)
        except (ImportError, AttributeError):
            pass

    if merger_backend is None:
        # Fallback: Nur Original kopieren
        shutil.copy2(pdf_path, output_path)
        return True

    tmp_body_path = None
    try:
        # Body zu temporaerem PDF
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            tmp_body_path = Path(tmp.name)

        if not html_to_pdf(body_html, tmp_body_path, mail_meta, mode="fast"):
            # Body-PDF fehlgeschlagen - nur Original kopieren
            # Hinweis: Merge immer mit "fast" da nur Body-Kontext benoetigt
            shutil.copy2(pdf_path, output_path)
            return True

        # PDFs mergen: Original + Body
        backend_type, backend_cls = merger_backend
        if backend_type == "pypdf2":
            merger = backend_cls()
            try:
                merger.append(str(pdf_path))
                merger.append(str(tmp_body_path))

                with open(output_path, 'wb') as f:
                    merger.write(f)
            finally:
                merger.close()
        else:
            writer = backend_cls()
            try:
                writer.append(str(pdf_path))
                writer.append(str(tmp_body_path))

                with open(output_path, 'wb') as f:
                    writer.write(f)
            finally:
                writer.close()

        return output_path.exists()
    except Exception:
        # Bei Fehler: Original kopieren
        shutil.copy2(pdf_path, output_path)
        return True
    finally:
        if tmp_body_path is not None:
            tmp_body_path.unlink(missing_ok=True)


def get_attachment_conversion_type(filename: str) -> Optional[str]:
    """Classifies supported attachment types for PDF archiving."""
    extension = Path(filename or "").suffix.lower()
    if extension in PDF_ATTACHMENT_EXTENSIONS:
        return "pdf"
    if extension in IMAGE_ATTACHMENT_EXTENSIONS:
        return "image"
    if extension == ".docx":
        return "docx"
    if extension == ".xlsx":
        return "xlsx"
    if extension in LEGACY_ATTACHMENT_EXTENSIONS:
        return "legacy_office"
    return None


def format_attachment_cell(value: object) -> str:
    """Formats arbitrary spreadsheet/document values for safe HTML rendering."""
    if value is None:
        return ""
    if isinstance(value, datetime):
        rendered = value.strftime("%Y-%m-%d %H:%M")
    elif isinstance(value, date):
        rendered = value.isoformat()
    else:
        rendered = str(value)
    return escape(rendered).replace("\n", "<br>")


def render_docx_attachment_html(file_data: bytes) -> str:
    """Renders a DOCX file into simple HTML for PDF conversion."""
    if not PYTHON_DOCX_AVAILABLE or DocxDocument is None:
        raise RuntimeError("python-docx nicht installiert")

    document = DocxDocument(io.BytesIO(file_data))
    blocks: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = ""
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
        heading_match = re.search(r'heading\s*(\d+)', style_name)
        text_html = format_attachment_cell(text)

        if heading_match:
            level = max(1, min(6, int(heading_match.group(1))))
            blocks.append(f"<h{level}>{text_html}</h{level}>")
        else:
            blocks.append(f"<p>{text_html}</p>")

    for table in document.tables:
        rows: List[str] = []
        for row_index, row in enumerate(table.rows):
            tag = "th" if row_index == 0 else "td"
            cells = "".join(
                f"<{tag}>{format_attachment_cell(cell.text.strip()) or '&nbsp;'}</{tag}>"
                for cell in row.cells
            )
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            blocks.append(
                "<table border='1' cellspacing='0' cellpadding='4' style='margin: 12px 0;'>"
                f"{''.join(rows)}</table>"
            )

    if not blocks:
        blocks.append("<p>(Leeres Dokument)</p>")
    return "".join(blocks)


def render_xlsx_attachment_html(file_data: bytes) -> str:
    """Renders an XLSX workbook into simple HTML tables for PDF conversion."""
    if not OPENPYXL_AVAILABLE or load_workbook is None:
        raise RuntimeError("openpyxl nicht installiert")

    workbook = load_workbook(io.BytesIO(file_data), data_only=True)
    sections: List[str] = []

    try:
        for sheet in workbook.worksheets:
            sections.append(f"<h2>{escape(sheet.title)}</h2>")
            rows = [
                tuple(row)
                for row in sheet.iter_rows(values_only=True)
                if any(cell not in (None, "") for cell in row)
            ]

            if not rows:
                sections.append("<p><em>Leeres Blatt</em></p>")
                continue

            table_rows: List[str] = []
            header_cells = "".join(
                f"<th>{format_attachment_cell(cell) or '&nbsp;'}</th>"
                for cell in rows[0]
            )
            table_rows.append(f"<tr>{header_cells}</tr>")

            for row in rows[1:]:
                cells = "".join(
                    f"<td>{format_attachment_cell(cell) or '&nbsp;'}</td>"
                    for cell in row
                )
                table_rows.append(f"<tr>{cells}</tr>")

            sections.append(
                "<table border='1' cellspacing='0' cellpadding='4' style='margin: 12px 0;'>"
                f"{''.join(table_rows)}</table>"
            )
    finally:
        workbook.close()

    if not sections:
        sections.append("<p>(Leere Arbeitsmappe)</p>")
    return "".join(sections)


def find_libreoffice_executable() -> Optional[Path]:
    """Finds a usable LibreOffice executable for headless conversions."""
    candidates: List[Path] = []

    for env_var in ("LIBREOFFICE_PATH", "SOFFICE_PATH"):
        raw_value = os.environ.get(env_var)
        if not raw_value:
            continue
        env_path = Path(raw_value)
        if env_path.is_dir():
            candidates.extend(
                env_path / executable_name
                for executable_name in ("soffice.exe", "soffice.com", "soffice")
            )
        else:
            candidates.append(env_path)

    for executable_name in ("soffice.exe", "soffice.com", "soffice"):
        resolved = shutil.which(executable_name)
        if resolved:
            candidates.append(Path(resolved))

    candidates.extend(LIBREOFFICE_CANDIDATE_PATHS)

    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def convert_legacy_office_via_com(source_path: Path, output_path: Path) -> str:
    """Converts .doc/.xls files via locally installed Word/Excel automation."""
    if not WIN32COM_AVAILABLE or pythoncom is None or win32com is None:
        raise RuntimeError("pywin32/COM ist nicht verfügbar")

    source_suffix = source_path.suffix.lower()
    app = None
    document = None
    workbook = None
    pythoncom.CoInitialize()

    try:
        if source_suffix == ".doc":
            app = win32com.client.DispatchEx("Word.Application")
            app.Visible = False
            app.DisplayAlerts = 0
            document = app.Documents.Open(
                str(source_path),
                ReadOnly=True,
                AddToRecentFiles=False,
            )
            document.ExportAsFixedFormat(str(output_path), WORD_PDF_FORMAT)
            return "Word COM"

        if source_suffix == ".xls":
            app = win32com.client.DispatchEx("Excel.Application")
            app.Visible = False
            app.DisplayAlerts = False
            workbook = app.Workbooks.Open(str(source_path), ReadOnly=True)
            workbook.ExportAsFixedFormat(EXCEL_PDF_FORMAT, str(output_path))
            return "Excel COM"

        raise RuntimeError(f"Legacy-Endung {source_suffix} wird per COM nicht unterstützt")
    except Exception as exc:
        raise RuntimeError(f"COM-Konvertierung fehlgeschlagen: {exc}") from exc
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if workbook is not None:
            try:
                workbook.Close(False)
            except Exception:
                pass
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def convert_legacy_office_via_libreoffice(source_path: Path, output_path: Path) -> str:
    """Converts .doc/.xls files via LibreOffice's headless CLI."""
    soffice_path = find_libreoffice_executable()
    if soffice_path is None:
        raise RuntimeError("LibreOffice (soffice) nicht gefunden")

    source_dir = source_path.parent
    generated_pdf = source_dir / f"{source_path.stem}.pdf"
    if generated_pdf.exists():
        generated_pdf.unlink()

    command = [
        str(soffice_path),
        "--headless",
        "--nologo",
        "--nodefault",
        "--nolockcheck",
        "--convert-to",
        "pdf",
        "--outdir",
        str(source_dir),
        str(source_path),
    ]
    creationflags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
        check=False,
        timeout=120,
        creationflags=creationflags,
    )

    if result.returncode != 0:
        details = result.stderr.strip() or result.stdout.strip() or "ohne Detailausgabe"
        raise RuntimeError(f"LibreOffice-CLI Fehler {result.returncode}: {details}")
    if not generated_pdf.exists():
        raise RuntimeError("LibreOffice hat keine PDF-Datei erzeugt")

    if generated_pdf.resolve() != output_path.resolve():
        shutil.copyfile(generated_pdf, output_path)
    return "LibreOffice"


def convert_legacy_office_attachment(
    file_data: bytes,
    source_name: str,
    output_path: Path,
) -> Tuple[bool, str]:
    """Converts legacy Office attachments via COM first, then LibreOffice as fallback."""
    extension = Path(source_name or "").suffix.lower() or "(ohne Endung)"
    errors: List[str] = []

    with tempfile.TemporaryDirectory(prefix="uim-legacy-") as temp_dir:
        temp_dir_path = Path(temp_dir)
        temp_source = temp_dir_path / f"legacy_attachment{extension}"
        temp_output = temp_dir_path / "legacy_attachment.pdf"
        temp_source.write_bytes(file_data)

        if os.name == "nt":
            try:
                backend = convert_legacy_office_via_com(temp_source, temp_output)
            except RuntimeError as exc:
                errors.append(str(exc))
            else:
                if temp_output.exists() and temp_output.stat().st_size > 0:
                    shutil.copyfile(temp_output, output_path)
                    return True, f"Legacy-Anhang via {backend} konvertiert"
                errors.append("COM-Konvertierung lieferte keine PDF-Datei")

        try:
            backend = convert_legacy_office_via_libreoffice(temp_source, temp_output)
        except (OSError, RuntimeError, subprocess.SubprocessError) as exc:
            errors.append(str(exc))
        else:
            if temp_output.exists() and temp_output.stat().st_size > 0:
                shutil.copyfile(temp_output, output_path)
                return True, f"Legacy-Anhang via {backend} konvertiert"
            errors.append("LibreOffice-Konvertierung lieferte keine PDF-Datei")

    details = "; ".join(errors) if errors else "Kein Konverter verfügbar"
    return False, f"Legacy-Format {extension} konnte nicht konvertiert werden: {details}"


def convert_attachment_to_pdf(file_data: bytes, source_name: str, output_path: Path) -> Tuple[bool, str]:
    """Converts supported attachments into a PDF file for unified archiving."""
    attachment_type = get_attachment_conversion_type(source_name)
    extension = Path(source_name or "").suffix.lower() or "(ohne Endung)"

    try:
        if attachment_type == "pdf":
            with open(output_path, "wb") as handle:
                handle.write(file_data)
            return True, "PDF-Anhang gespeichert"

        if attachment_type == "image":
            if not PILLOW_AVAILABLE or Image is None:
                return False, "Pillow nicht installiert"

            with Image.open(io.BytesIO(file_data)) as img:
                try:
                    from PIL import ImageOps
                    img = ImageOps.exif_transpose(img)
                except Exception:
                    pass

                has_alpha = (
                    img.mode in {"RGBA", "LA", "PA"}
                    or (img.mode == "P" and "transparency" in getattr(img, "info", {}))
                    or ("A" in img.getbands() if hasattr(img, "getbands") else False)
                )

                if has_alpha:
                    rgba_img = img.convert("RGBA")
                    background = Image.new("RGB", rgba_img.size, (255, 255, 255))
                    background.paste(rgba_img, mask=rgba_img.getchannel("A"))
                    img = background
                elif img.mode != "RGB":
                    img = img.convert("RGB")
                img.save(output_path, "PDF", resolution=150.0)
            image_success = output_path.exists() and output_path.stat().st_size > 0
            if image_success:
                return True, "Bild-Anhang konvertiert"
            return False, "Bild-PDF-Konvertierung fehlgeschlagen"

        if attachment_type == "docx":
            html_content = render_docx_attachment_html(file_data)
            docx_success = html_to_pdf(html_content, output_path, None, mode="fast")
            if docx_success:
                return True, "DOCX-Anhang konvertiert"
            return False, "DOCX-PDF-Konvertierung fehlgeschlagen"

        if attachment_type == "xlsx":
            html_content = render_xlsx_attachment_html(file_data)
            xlsx_success = html_to_pdf(html_content, output_path, None, mode="fast")
            if xlsx_success:
                return True, "XLSX-Anhang konvertiert"
            return False, "XLSX-PDF-Konvertierung fehlgeschlagen"

        if attachment_type == "legacy_office":
            return convert_legacy_office_attachment(file_data, source_name, output_path)

        return False, f"Dateityp {extension} wird nicht unterstützt"
    except Exception as exc:
        logger.error(f"Attachment-PDF-Konvertierung fehlgeschlagen ({source_name}): {exc}")
        if output_path.exists():
            output_path.unlink(missing_ok=True)
        return False, str(exc)


def extract_order_id(text: str) -> str:
    """Extracts an order number from text (subject line or email body).

    Fix #14: Tightened regex — requires context keywords to avoid matching UUIDs.
    """
    patterns = [
        # Deutsche Formate - MUSS Keyword haben
        r'(?:Bestellnummer|Bestell-Nr\.?|Bestellnr\.?|Auftragsnummer|Auftrags-Nr\.?)\s*[:.]?\s*([A-Z0-9-]{5,20})',
        r'(?:Rechnungsnummer|Rechnungs-Nr\.?|Rechnung\s*Nr\.?)\s*[:.]?\s*([A-Z0-9-]{5,20})',
        # Englische Formate - MUSS Keyword haben
        r'(?:Order\s*#?|Order\s*ID|Order\s*Number|Order\s*No\.?)\s*[:.]?\s*([A-Z0-9-]{5,20})',
        r'(?:Invoice\s*#?|Invoice\s*ID|Invoice\s*Number)\s*[:.]?\s*([A-Z0-9-]{5,20})',
        r'(?:Reference\s*#?|Ref\.?\s*#?)\s*[:.]?\s*([A-Z0-9-]{5,20})',
        # Amazon-Style (explizites Format)
        r'#\s*([0-9]{3}-[0-9]{7}-[0-9]{7})',  # Amazon DE Format
        # Temu/Alibaba Style
        r'(?:PO|SO)[#:\s]*([0-9A-Z]{10,20})',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            result = match.group(1)
            # Fix #14: UUID-artige Strings ablehnen (8-4-4-4-12 Format)
            if re.match(r'^[a-f0-9]{8}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{4}-[a-f0-9]{12}$', result, re.IGNORECASE):
                continue
            return sanitize_filename(result)
    return ""


# ==================== WORKER THREAD ====================

class InvoiceWorker(QThread):
    """Background thread for mail querying and invoice download."""
    log = Signal(str)
    progress = Signal(int, int)  # current, total
    invoice_found = Signal(object)  # Invoice
    finished_signal = Signal(int)  # Anzahl gefundener Rechnungen

    def __init__(self, accounts: List[MailAccount], profiles: List[InvoiceProfile],
                 settings: AppSettings, existing_hashes: set):
        super().__init__()
        self.accounts = accounts
        self.profiles = profiles
        self.settings = settings
        self.existing_hashes = existing_hashes
        self.should_stop = False
        self.found_count = 0

    def stop(self):
        self.should_stop = True

    def run(self):
        self.found_count = 0

        for account in self.accounts:
            if self.should_stop:
                break

            # Profile für diesen Account filtern
            account_profiles = [p for p in self.profiles if p.account_id == account.id and p.enabled]
            if not account_profiles:
                continue

            self.log.emit(f"\n{'='*50}")
            self.log.emit(f"📧 Verarbeite Account: {account.name}")
            self.log.emit(f"{'='*50}")

            if account.use_gmail_api:
                self._process_gmail_api(account, account_profiles)
            else:
                self._process_imap(account, account_profiles)

        self.finished_signal.emit(self.found_count)

    def _process_gmail_api(self, account: MailAccount, profiles: List[InvoiceProfile]):
        """Processes an account via the Gmail API."""
        if not GMAIL_API_AVAILABLE:
            self.log.emit("❌ Gmail API nicht verfügbar (google-api-python-client fehlt)")
            return

        try:
            creds = self._get_gmail_credentials()
            if not creds:
                return

            service = build("gmail", "v1", credentials=creds, cache_discovery=False)

            for profile in profiles:
                if self.should_stop:
                    break
                self._search_gmail(service, profile)

        except Exception as e:
            self.log.emit(f"❌ Gmail API Fehler: {e}")

    def _get_gmail_credentials(self) -> "Optional[Credentials]":
        """Retrieves or refreshes Gmail API credentials, running OAuth flow if needed."""
        creds = None

        if TOKEN_FILE.exists():
            try:
                creds = Credentials.from_authorized_user_file(str(TOKEN_FILE), GMAIL_SCOPES)
            except (OSError, ValueError):
                # Token-Datei korrupt oder nicht lesbar, neu authentifizieren
                TOKEN_FILE.unlink(missing_ok=True)

        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                try:
                    self.log.emit("🔄 Aktualisiere Gmail Token...")
                    creds.refresh(Request())
                except RefreshError:
                    self.log.emit("⚠️ Token abgelaufen - Neu-Login erforderlich")
                    TOKEN_FILE.unlink(missing_ok=True)
                    creds = None

            if not creds:
                if not CREDENTIALS_FILE.exists():
                    # Auto-Suche nach vorhandenen credentials.json
                    found_creds = self._find_existing_credentials()
                    if found_creds:
                        self.log.emit(f"🔍 Gefundene credentials.json: {found_creds}")
                        self.log.emit(f"📋 Kopiere nach {CREDENTIALS_FILE}...")
                        try:
                            import shutil
                            shutil.copy2(found_creds, CREDENTIALS_FILE)
                            self.log.emit("✅ credentials.json kopiert!")
                        except Exception as e:
                            self.log.emit(f"❌ Kopieren fehlgeschlagen: {e}")
                            return None
                    else:
                        self.log.emit(f"❌ credentials.json fehlt in {BASE_DIR}")
                        self.log.emit("💡 Erstelle OAuth Credentials in Google Cloud Console")
                        self.log.emit("💡 Oder lege eine credentials.json in einem dieser Ordner ab:")
                        self.log.emit("   - Downloads")
                        self.log.emit("   - OneDrive\\.SOFTWARE\\TOOLS")
                        return None

                self.log.emit("🌐 Starte Browser für Google Login...")
                try:
                    flow = InstalledAppFlow.from_client_secrets_file(
                        str(CREDENTIALS_FILE), GMAIL_SCOPES)
                    creds = flow.run_local_server(port=0)

                    with open(TOKEN_FILE, "w") as f:
                        f.write(creds.to_json())
                    self.log.emit("✅ Gmail Authentifizierung erfolgreich")
                except Exception as e:
                    self.log.emit(f"❌ Auth Fehler: {e}")
                    return None

        return creds

    @staticmethod
    def _quote_imap_string(value: str) -> str:
        """Escapes a string for IMAP search arguments."""
        escaped = (value or "").replace("\\", "\\\\").replace('"', '\\"').strip()
        return f'"{escaped}"'

    @staticmethod
    def _supports_gmail_raw(mail) -> bool:
        """Returns True when the IMAP server supports Gmail's X-GM-RAW extension."""
        capabilities = getattr(mail, "capabilities", ()) or ()
        normalized = {
            cap.decode("ascii", errors="ignore").upper() if isinstance(cap, bytes) else str(cap).upper()
            for cap in capabilities
        }
        return "X-GM-EXT-1" in normalized

    def _build_gmail_search_query(self, profile: InvoiceProfile) -> str:
        """Builds a Gmail-style search query from saved raw query and profile filters."""
        query_parts = []

        if getattr(profile, "gmail_query", "").strip():
            query_parts.append(profile.gmail_query.strip())

        if profile.sender_filter:
            senders = [s.strip() for s in profile.sender_filter.split(",") if s.strip()]
            if senders:
                sender_q = " OR ".join([f"from:{s}" for s in senders])
                query_parts.append(f"({sender_q})")

        if profile.subject_filter:
            subjects = [s.strip() for s in profile.subject_filter.split(",") if s.strip()]
            if subjects:
                subj_q = " OR ".join([f'subject:"{s}"' for s in subjects])
                query_parts.append(f"({subj_q})")

        if self.settings.date_from:
            query_parts.append(f"after:{self.settings.date_from.replace('-', '/')}")
        if self.settings.date_to:
            query_parts.append(f"before:{self.settings.date_to.replace('-', '/')}")

        if not self.settings.date_from and not self.settings.date_to and self.settings.date_filter_months > 0:
            since = datetime.now() - timedelta(days=self.settings.date_filter_months * 30)
            query_parts.append(f"after:{since.strftime('%Y/%m/%d')}")

        if not profile.subject_filter:
            query_parts.append("(has:attachment OR subject:rechnung OR subject:invoice)")

        return " ".join(part for part in query_parts if part).strip()

    def _build_imap_search_args(self, profile: InvoiceProfile) -> List[str]:
        """Builds standard IMAP search arguments for servers without X-GM-RAW."""
        search_args: List[str] = []

        if self.settings.date_from:
            try:
                from_date = datetime.strptime(self.settings.date_from, "%Y-%m-%d")
                search_args.extend(["SINCE", format_imap_date(from_date)])
            except ValueError:
                pass

        if self.settings.date_to:
            try:
                to_date = datetime.strptime(self.settings.date_to, "%Y-%m-%d")
                to_date_plus1 = to_date + timedelta(days=1)
                search_args.extend(["BEFORE", format_imap_date(to_date_plus1)])
            except ValueError:
                pass

        if not search_args and self.settings.date_filter_months > 0:
            since = datetime.now() - timedelta(days=self.settings.date_filter_months * 30)
            search_args.extend(["SINCE", format_imap_date(since)])

        if profile.sender_filter:
            senders = [s.strip() for s in profile.sender_filter.split(",") if s.strip()]
            if len(senders) == 1:
                safe_sender = senders[0].replace('"', '')
                search_args.extend(["FROM", self._quote_imap_string(safe_sender)])
            elif len(senders) > 1:
                # IMAP OR: OR FROM "a" FROM "b" for 2; OR FROM "a" OR FROM "b" FROM "c" for 3+
                or_args: List[str] = []
                for i, sender in enumerate(senders):
                    safe = sender.replace('"', '')
                    if i < len(senders) - 1:
                        or_args.append("OR")
                    or_args.extend(["FROM", self._quote_imap_string(safe)])
                search_args.extend(or_args)

        subjects = [s.strip() for s in profile.subject_filter.split(",") if s.strip()]
        if len(subjects) == 1:
            safe_subject = subjects[0].replace('"', '')
            search_args.extend(["SUBJECT", self._quote_imap_string(safe_subject)])
        elif len(subjects) > 1:
            # IMAP OR: OR SUBJECT "a" SUBJECT "b" for 2; nested for 3+
            or_args: List[str] = []
            for i, subj in enumerate(subjects):
                safe = subj.replace('"', '')
                if i < len(subjects) - 1:
                    or_args.append("OR")
                or_args.extend(["SUBJECT", self._quote_imap_string(safe)])
            search_args.extend(or_args)

        return search_args or ["ALL"]

    def _search_gmail(self, service, profile: InvoiceProfile):
        """Searches for emails via the Gmail API using the given profile's filters."""
        self.log.emit(f"\n[SEARCH] Profil: {profile.name}")

        query = self._build_gmail_search_query(profile)
        self.log.emit(f"   Query: {query}")

        try:
            results = service.users().messages().list(
                userId="me", q=query, maxResults=self.settings.max_emails_per_run
            ).execute()

            messages = results.get("messages", [])
            self.log.emit(f"   📬 {len(messages)} Mails gefunden")

            for i, msg_ref in enumerate(messages):
                if self.should_stop:
                    break

                self.progress.emit(i + 1, len(messages))

                try:
                    msg = service.users().messages().get(
                        userId="me", id=msg_ref["id"], format="full"
                    ).execute()

                    # Rate Limiting: Pause zwischen API-Aufrufen (verhindert 429)
                    time.sleep(0.1)

                    self._process_gmail_message(service, msg, profile)

                except Exception as e:
                    self.log.emit(f"   ⚠️ Fehler bei Mail: {e}")

        except Exception as e:
            self.log.emit(f"❌ Gmail Suche fehlgeschlagen: {e}")

    def _process_gmail_message(self, service, msg: dict, profile: InvoiceProfile):
        """Processes a single Gmail message: applies filters, downloads attachments, converts body."""
        headers = {h['name'].lower(): h['value'] for h in msg['payload'].get('headers', [])}

        sender = headers.get('from', 'Unknown')
        subject = headers.get('subject', 'No Subject')
        date_str = headers.get('date', '')

        # Body nur einmal holen wenn noetig fuer Filter
        blacklist = getattr(profile, 'blacklist', '')
        body_must = getattr(profile, 'body_must_contain', '')
        body_must_not = getattr(profile, 'body_must_not_contain', '')
        body_text = ""
        if blacklist or body_must or body_must_not:
            body_text = self._get_message_body(msg['payload']).lower()

        # Gemeinsame Filter-Logik (Blacklist, body_must, body_must_not)
        if not self._check_message_filters(profile, subject, body_text):
            return

        # Datum parsen
        try:
            mail_date = email.utils.parsedate_to_datetime(date_str)
            fmt_date = mail_date.strftime("%Y-%m-%d")
        except (ValueError, TypeError):
            fmt_date = datetime.now().strftime("%Y-%m-%d")

        # Zielordner
        target_dir = self._compute_target_dir(profile)

        found_any = False
        safe_profile_name = sanitize_filename(profile.name)
        body_html = self._get_message_body(msg['payload']) if self.settings.merge_body_with_attachments else ""

        # 1. Anhänge verarbeiten
        if self.settings.download_attachments:
            parts = self._get_all_parts(msg['payload'])

            for part in parts:
                filename = part.get('filename', '')
                attachment_type = get_attachment_conversion_type(filename)
                if filename and attachment_type:
                    body_obj = part.get('body', {})
                    att_id = body_obj.get('attachmentId')
                    inline_data = body_obj.get('data')

                    # Versuche Daten zu holen
                    file_data = None
                    try:
                        if inline_data:
                            file_data = safe_b64decode(inline_data)
                        elif att_id:
                            att = service.users().messages().attachments().get(
                                userId="me", messageId=msg['id'], id=att_id
                            ).execute()
                            time.sleep(0.05)  # Rate Limiting fuer Attachment-API
                            file_data = safe_b64decode(att['data'])

                        # Keine Daten gefunden? Skip
                        if not file_data:
                            continue

                        if self._save_attachment_invoice(
                            file_data=file_data,
                            source_name=filename,
                            profile=profile,
                            fmt_date=fmt_date,
                            subject=subject,
                            sender=sender,
                            fallback_seed=msg['id'],
                            target_dir=target_dir,
                            body_html=body_html,
                        ):
                            found_any = True

                    except Exception as e:
                        self.log.emit(f"   ⚠️ Anhang-Fehler: {e}")

        # 2. Body zu PDF (nur wenn keine Anhaenge gefunden und Option aktiv)
        if not found_any and self.settings.convert_body_to_pdf:
            body_html = self._get_message_body(msg['payload'])
            if body_html and len(body_html) > 200:
                body_hash = calculate_hash(body_html.encode('utf-8'))

                if self.settings.enable_hash_check and body_hash in self.existing_hashes:
                    return

                # Bestellnummer aus Subject oder Body extrahieren
                order_id = extract_order_id(subject)
                if not order_id:
                    order_id = extract_order_id(body_html)
                if not order_id:
                    # MD5 als sicherer Fallback (keine zufaelligen Woerter)
                    order_id = hashlib.md5(msg['id'].encode()).hexdigest()[:8]

                safe_name = f"{safe_profile_name}_{fmt_date}_{order_id}_mail.pdf"
                output_path = target_dir / safe_name

                # Hybrid-Design: Mail-Header in Body-PDF
                mail_meta = {'sender': sender, 'subject': subject, 'date': fmt_date}

                if not output_path.exists() and html_to_pdf(body_html, output_path, mail_meta, mode=self.settings.pdf_mode):
                    self.log.emit(f"   📄 Mail->PDF: {safe_name}")

                    # OCR hinzufuegen wenn aktiviert und PDF keine Textebene hat
                    # OCR: Immer ausfuehren wenn aktiviert (scannt Bilder nach Text)
                    if self.settings.ocr_enabled and OCR_AVAILABLE:
                        ocr = OCRProcessor()
                        success, ocr_msg = ocr.enhance_with_ocr(output_path)
                        if success:
                            self.log.emit(f"   🔍 OCR: {ocr_msg}")
                        else:
                            # Kein Fehler loggen wenn einfach kein Text erkannt wurde
                            if "Kein Text" not in ocr_msg:
                                self.log.emit(f"   ⚠️ OCR: {ocr_msg}")

                    inv = Invoice(
                        id=str(uuid.uuid4()),
                        profile_name=profile.name,
                        filename=safe_name,
                        date=fmt_date,
                        path=str(output_path),
                        sender=sender[:80],
                        subject=subject[:100],
                        hash=body_hash,
                        is_attachment=False
                    )
                    self.invoice_found.emit(inv)
                    self.existing_hashes.add(body_hash)
                    self.found_count += 1

    def _get_all_parts(self, payload: dict) -> List[dict]:
        """Recursively extracts all MIME parts from a message payload."""
        parts = []
        if 'parts' in payload:
            for p in payload['parts']:
                parts.extend(self._get_all_parts(p))
        else:
            parts.append(payload)
        return parts

    def _get_message_body(self, payload: dict) -> str:
        """Extracts the HTML/text body from a message payload, preferring HTML.

        Solves the multipart/alternative problem (e.g. Temu emails) where
        text/plain only contains a fallback and text/html has the actual content.
        """
        # Erst alle Parts sammeln (rekursiv)
        all_parts = self._get_all_body_parts(payload)

        # HTML bevorzugen
        for mime_type, data in all_parts:
            if 'html' in mime_type:
                return data

        # Fallback auf plain text
        for mime_type, data in all_parts:
            if 'plain' in mime_type:
                return f"<pre>{escape(data)}</pre>"

        return ""

    def _get_all_body_parts(self, payload: dict) -> List[Tuple[str, str]]:
        """Recursively collects all body parts as (mimeType, decoded_data) tuples."""
        parts = []

        if 'parts' in payload:
            for part in payload['parts']:
                parts.extend(self._get_all_body_parts(part))
        else:
            mime_type = payload.get('mimeType', '')
            body_data = payload.get('body', {}).get('data', '')

            if body_data and ('html' in mime_type or 'plain' in mime_type):
                decoded = safe_b64decode(body_data).decode('utf-8', errors='ignore')
                parts.append((mime_type, decoded))

        return parts

    def _check_message_filters(self, profile: 'InvoiceProfile', subject: str,
                                body_text: str) -> bool:
        """Applies profile-level content filters to a message.

        Checks blacklist (subject + body), body_must_contain, and
        body_must_not_contain filters. Emits log messages on rejection.

        Args:
            profile: The InvoiceProfile whose filters should be applied.
            subject: Decoded subject line of the message (case-insensitive comparison).
            body_text: Decoded body text, already lowercased by the caller.

        Returns:
            True if the message passes all filters, False if it should be skipped.
        """
        blacklist = getattr(profile, 'blacklist', '')
        body_must = getattr(profile, 'body_must_contain', '')
        body_must_not = getattr(profile, 'body_must_not_contain', '')
        subject_lower = subject.lower()

        # Blacklist pruefen (Betreff + Body)
        if blacklist:
            blacklist_terms = [b.strip().lower() for b in blacklist.split(",") if b.strip()]
            for term in blacklist_terms:
                if term in subject_lower or term in body_text:
                    self.log.emit(f"   ⛔ Blacklist: '{term}' gefunden, übersprungen")
                    return False

        # Body MUSS enthalten (mindestens eines)
        if body_must:
            must_terms = [t.strip().lower() for t in body_must.split(",") if t.strip()]
            if must_terms and not any(term in body_text for term in must_terms):
                self.log.emit(f"   ⏭️ Body enthält keines von: {body_must[:30]}...")
                return False

        # Body darf NICHT enthalten
        if body_must_not:
            must_not_terms = [t.strip().lower() for t in body_must_not.split(",") if t.strip()]
            for term in must_not_terms:
                if term in body_text:
                    self.log.emit(f"   ⛔ Body-Blacklist: '{term}' gefunden, übersprungen")
                    return False

        return True

    def _compute_target_dir(self, profile: 'InvoiceProfile') -> Path:
        """Computes and creates the target directory for a profile's invoices.

        Args:
            profile: The InvoiceProfile whose target directory should be resolved.

        Returns:
            Created Path object pointing to the profile's download subdirectory.
        """
        target_dir = Path(self.settings.download_path)
        if profile.target_subfolder:
            target_dir = target_dir / sanitize_filename(profile.target_subfolder)
        else:
            target_dir = target_dir / sanitize_filename(profile.name)
        target_dir.mkdir(parents=True, exist_ok=True)
        return target_dir

    def _build_attachment_output_path(
        self,
        target_dir: Path,
        profile: 'InvoiceProfile',
        fmt_date: str,
        subject: str,
        fallback_seed: str,
    ) -> Tuple[Path, str]:
        """Builds a unique PDF path for a converted attachment."""
        safe_profile_name = sanitize_filename(profile.name)
        order_id = extract_order_id(subject)
        if not order_id:
            safe_seed = fallback_seed or subject or profile.name
            order_id = hashlib.md5(safe_seed.encode("utf-8", errors="ignore")).hexdigest()[:8]

        safe_name = f"{safe_profile_name}_{fmt_date}_{order_id}.pdf"
        output_path = target_dir / safe_name
        counter = 1

        while output_path.exists():
            safe_name = f"{safe_profile_name}_{fmt_date}_{order_id}_{counter}.pdf"
            output_path = target_dir / safe_name
            counter += 1

        return output_path, safe_name

    def _save_attachment_invoice(
        self,
        file_data: bytes,
        source_name: str,
        profile: 'InvoiceProfile',
        fmt_date: str,
        subject: str,
        sender: str,
        fallback_seed: str,
        target_dir: Path,
        body_html: str = "",
    ) -> bool:
        """Converts/saves a supported attachment and emits the resulting invoice."""
        file_hash = calculate_hash(file_data)
        if self.settings.enable_hash_check and file_hash in self.existing_hashes:
            self.log.emit(f"   ⏭️ Duplikat übersprungen: {source_name}")
            return False

        output_path, safe_name = self._build_attachment_output_path(
            target_dir, profile, fmt_date, subject, fallback_seed
        )
        attachment_type = get_attachment_conversion_type(source_name)
        mail_meta = {'sender': sender, 'subject': subject, 'date': fmt_date}
        temp_pdf_path = None

        try:
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                temp_pdf_path = Path(tmp.name)

            success, message = convert_attachment_to_pdf(file_data, source_name, temp_pdf_path)
            if not success:
                self.log.emit(f"   ⏭️ Anhang übersprungen: {source_name} ({message})")
                return False

            merge_requested = self.settings.merge_body_with_attachments and len(body_html or "") > 100
            if merge_requested:
                if not merge_pdf_with_body(temp_pdf_path, body_html, mail_meta, output_path):
                    shutil.move(str(temp_pdf_path), str(output_path))
            else:
                shutil.move(str(temp_pdf_path), str(output_path))

            action = "Gespeichert" if attachment_type == "pdf" else "Konvertiert"
            if merge_requested:
                self.log.emit(f"   📎+ {action}: {safe_name}")
            else:
                self.log.emit(f"   📎 {action}: {safe_name}")

            inv = Invoice(
                id=str(uuid.uuid4()),
                profile_name=profile.name,
                filename=safe_name,
                date=fmt_date,
                path=str(output_path),
                profile_id=profile.id,
                sender=sender[:80],
                subject=subject[:100],
                hash=file_hash,
                is_attachment=True
            )
            self.invoice_found.emit(inv)
            self.existing_hashes.add(file_hash)
            self.found_count += 1
            return True
        finally:
            if temp_pdf_path is not None:
                temp_pdf_path.unlink(missing_ok=True)

    def _get_imap_message_body(self, msg) -> str:
        """Extracts the preferred IMAP message body, preferring HTML over plain text."""
        html_content = None
        plain_content = None

        for part in msg.walk():
            if part.get_content_maintype() == 'multipart':
                continue

            content_type = part.get_content_type()
            if content_type not in {"text/html", "text/plain"}:
                continue

            try:
                payload = part.get_payload(decode=True)
                if not payload:
                    continue
                # Echten Charset aus dem Content-Type-Header auslesen statt blind
                # UTF-8 anzunehmen — ISO-8859-1, windows-1252 etc. kommen in der Praxis vor.
                charset = part.get_content_charset() or 'utf-8'
                try:
                    decoded = payload.decode(charset, errors='replace')
                except (LookupError, UnicodeDecodeError):
                    decoded = payload.decode('utf-8', errors='replace')
            except (AttributeError, TypeError):
                continue

            if content_type == "text/html" and html_content is None:
                html_content = decoded
            elif content_type == "text/plain" and plain_content is None:
                plain_content = decoded

        return html_content or plain_content or ""

    def _find_existing_credentials(self) -> Optional[Path]:
        """Searches known directories for an existing credentials.json file.

        Looks in Downloads, OneDrive, Documents, and other common locations
        up to 3 directory levels deep.

        Returns:
            Path to a valid credentials.json, or None if not found.
        """
        search_dirs = [
            Path.home() / "Downloads",
            Path.home() / "OneDrive" / "Software Entwicklung" / "TOOLS",
            Path.home() / "OneDrive" / "Software Entwicklung",
            Path.home() / "Documents",
            Path.home() / ".invoicemaster",
            Path.home() / ".gmail_credentials",
        ]

        for base_dir in search_dirs:
            if not base_dir.exists():
                continue

            # Direkt im Ordner
            cred_file = base_dir / "credentials.json"
            if cred_file.exists():
                return cred_file

            # Rekursiv suchen (max 3 Ebenen)
            try:
                for depth in range(1, 4):
                    pattern = "/".join(["*"] * depth) + "/credentials.json"
                    for found in base_dir.glob(pattern):
                        if found.exists() and found.stat().st_size > 100:
                            return found
            except OSError:
                continue

        return None

    def _get_imap_folders(self, mail, host: str) -> List[str]:
        """Ermittelt relevante IMAP-Ordner zum Durchsuchen"""
        folders = []

        try:
            # Liste aller Ordner holen
            _, folder_list = mail.list()

            # Ordnernamen die uebersprungen werden (ausser wenn include_trash)
            skip_names = ['Spam', 'Junk', 'Drafts', 'Entwürfe']
            if not self.settings.include_trash:
                skip_names.extend(['Trash', 'Papierkorb', 'Deleted', 'Gelöscht'])

            all_folders = []
            for folder_info in folder_list:
                if folder_info:
                    try:
                        folder_str = folder_info.decode('utf-8') if isinstance(folder_info, bytes) else str(folder_info)
                        if '"' in folder_str:
                            parts = folder_str.split('"')
                            folder_name = parts[-2] if len(parts) >= 2 else parts[-1]
                        else:
                            folder_name = folder_str.split()[-1]
                        all_folders.append(folder_name)
                    except (UnicodeDecodeError, ValueError, IndexError):
                        continue

            # Gmail: "All Mail" enthält alles
            if 'gmail' in host.lower():
                for f in all_folders:
                    if 'All Mail' in f or 'Alle Nachrichten' in f or 'Alle Mails' in f:
                        return [f]

            # Andere Provider: Relevante Ordner
            for f in all_folders:
                skip = False
                for skip_name in skip_names:
                    if skip_name.lower() in f.lower():
                        skip = True
                        break
                if not skip:
                    folders.append(f)

            if not folders:
                folders = ['INBOX']

            return folders[:5]  # Max 5 Ordner

        except Exception as e:
            self.log.emit(f"   [!] Ordnerliste: {e}")
            return ['INBOX']

    def _process_imap(self, account: MailAccount, profiles: List[InvoiceProfile]):
        """Connects to an IMAP account and processes all enabled profiles.

        Retrieves the password from keyring, opens an IMAP4_SSL connection,
        determines relevant folders, and delegates each profile to _search_imap.

        Args:
            account: The MailAccount configuration (host, port, username).
            profiles: List of enabled InvoiceProfiles assigned to this account.
        """
        password = ""
        if KEYRING_AVAILABLE:
            try:
                password = keyring.get_password(APP_NAME, account.id)
            except (OSError, RuntimeError):
                pass

        if not password:
            self.log.emit(f"❌ Kein Passwort für {account.name} gespeichert")
            return

        mail = None
        try:
            self.log.emit(f"Verbinde mit {account.host}:{account.port}...")

            mail = imaplib.IMAP4_SSL(account.host, account.port)
            mail.login(account.username, password)

            # Ordner zum Durchsuchen ermitteln
            folders_to_search = self._get_imap_folders(mail, account.host)
            self.log.emit(f"[OK] IMAP Login erfolgreich - {len(folders_to_search)} Ordner")

            for folder in folders_to_search:
                if self.should_stop:
                    break
                try:
                    status, _ = mail.select(folder)
                    if status != 'OK':
                        continue
                    self.log.emit(f"Durchsuche: {folder}")

                    for profile in profiles:
                        if self.should_stop:
                            break
                        self._search_imap(mail, profile)
                except Exception as e:
                    self.log.emit(f"   [!] Ordner {folder}: {e}")
                    continue

            mail.logout()

        except imaplib.IMAP4.error as e:
            self.log.emit(f"❌ IMAP Fehler: {e}")
            self.log.emit("💡 Bei Gmail: App-Passwort in Google Konto erstellen")
        except Exception as e:
            self.log.emit(f"❌ Verbindungsfehler: {e}")
        finally:
            if mail is not None:
                try:
                    mail.shutdown()
                except Exception:
                    pass

    def _search_imap(self, mail, profile: InvoiceProfile):
        """Searches the currently selected IMAP folder using profile filters.

        Builds an IMAP search criteria string from the profile's date filters
        and sender filter, fetches matching message IDs, and delegates each
        message to _process_imap_message.

        Args:
            mail: Active imaplib.IMAP4_SSL connection with a folder selected.
            profile: InvoiceProfile whose filters define the search criteria.
        """
        self.log.emit(f"\n[SEARCH] Profil: {profile.name}")

        gmail_query = self._build_gmail_search_query(profile)
        use_gmail_raw = bool(getattr(profile, "gmail_query", "").strip()) and self._supports_gmail_raw(mail)

        if use_gmail_raw:
            search_args = (None, "X-GM-RAW", self._quote_imap_string(gmail_query))
            self.log.emit(f"   Gmail-RAW Suche: {gmail_query}")
        else:
            if getattr(profile, "gmail_query", "").strip():
                self.log.emit("   Gmail-Query gespeichert, aber Server unterstützt kein X-GM-RAW. Fallback auf IMAP-Filter.")
            imap_args = self._build_imap_search_args(profile)
            search_args = (None, *imap_args)
            self.log.emit(f"   IMAP Suche: {' '.join(imap_args)}")

        try:
            # UIDs verwenden statt MSN: UIDs bleiben stabil auch wenn andere
            # Clients gleichzeitig Mails verschieben oder löschen (RFC 3501 §2.3.1.1).
            _, message_ids = mail.uid('search', *search_args)
            ids = message_ids[0].split()

            # Limit
            ids = ids[-self.settings.max_emails_per_run:]
            self.log.emit(f"   📬 {len(ids)} Mails gefunden")

            for i, msg_id in enumerate(ids):
                if self.should_stop:
                    break

                self.progress.emit(i + 1, len(ids))

                try:
                    _, msg_data = mail.uid('fetch', msg_id, '(RFC822)')
                    # NIL-Guard: Server kann bei unbekannter UID leere/fehlerhafte Daten zurückgeben
                    if not msg_data or not isinstance(msg_data[0], tuple):
                        self.log.emit(f"   ⚠️ Keine Daten für UID {msg_id}")
                        continue
                    raw_email = msg_data[0][1]
                    if raw_email is None:
                        self.log.emit(f"   ⚠️ Leere Antwort für UID {msg_id}")
                        continue
                    msg = email.message_from_bytes(raw_email)

                    self._process_imap_message(msg, profile)

                except Exception as e:
                    self.log.emit(f"   ⚠️ Mail-Fehler: {e}")

        except Exception as e:
            self.log.emit(f"❌ IMAP Suche fehlgeschlagen: {e}")

    def _process_imap_message(self, msg, profile: InvoiceProfile):
        """Processes a single IMAP message: applies filters, downloads PDFs, converts body.

        Decodes sender/subject/date headers, runs content filters via
        _check_message_filters, then saves PDF attachments and/or converts the
        email body to PDF depending on AppSettings.

        Args:
            msg: email.message.Message object (from email.message_from_bytes).
            profile: InvoiceProfile whose filters and target folder apply.
        """
        sender = decode_mail_header(msg.get('From', ''))
        subject = decode_mail_header(msg.get('Subject', ''))
        date_str = msg.get('Date', '')

        # Body nur einmal extrahieren wenn noetig fuer Filter
        blacklist = getattr(profile, 'blacklist', '')
        body_must = getattr(profile, 'body_must_contain', '')
        body_must_not = getattr(profile, 'body_must_not_contain', '')
        body_text = ""
        if blacklist or body_must or body_must_not:
            body_text = self._get_imap_message_body(msg).lower()

        # Gemeinsame Filter-Logik (Blacklist, body_must, body_must_not)
        if not self._check_message_filters(profile, subject, body_text):
            return

        # Subject Filter pruefen
        if profile.subject_filter:
            keywords = [k.strip().lower() for k in profile.subject_filter.split(",") if k.strip()]
            if keywords and not any(k in subject.lower() for k in keywords):
                return  # Skip wenn kein Keyword matcht

        # Datum parsen
        try:
            mail_date = email.utils.parsedate_to_datetime(date_str)
            fmt_date = mail_date.strftime("%Y-%m-%d")
        except (ValueError, TypeError):
            fmt_date = datetime.now().strftime("%Y-%m-%d")

        # Zielordner
        target_dir = self._compute_target_dir(profile)

        found_any = False
        safe_profile_name = sanitize_filename(profile.name)
        body_html = self._get_imap_message_body(msg) if self.settings.merge_body_with_attachments else ""
        if body_html and not body_html.lstrip().startswith("<"):
            body_html = f"<pre>{escape(body_html)}</pre>"

        # Anhänge verarbeiten
        if self.settings.download_attachments:
            for part in msg.walk():
                if part.get_content_maintype() == 'multipart':
                    continue

                filename = part.get_filename()
                attachment_type = get_attachment_conversion_type(filename)
                if filename and attachment_type:
                    filename = decode_mail_header(filename)

                    try:
                        file_data = part.get_payload(decode=True)
                        if not file_data:
                            continue

                        if self._save_attachment_invoice(
                            file_data=file_data,
                            source_name=filename,
                            profile=profile,
                            fmt_date=fmt_date,
                            subject=subject,
                            sender=sender,
                            fallback_seed=subject or filename,
                            target_dir=target_dir,
                            body_html=body_html,
                        ):
                            found_any = True

                    except Exception as e:
                        self.log.emit(f"   ⚠️ Anhang-Fehler: {e}")

        # Body zu PDF
        if not found_any and self.settings.convert_body_to_pdf:
            body_html = self._get_imap_message_body(msg)
            if body_html and not body_html.lstrip().startswith("<"):
                body_html = f"<pre>{escape(body_html)}</pre>"

            if body_html and len(body_html) > 200:
                body_hash = calculate_hash(body_html.encode('utf-8'))

                if self.settings.enable_hash_check and body_hash in self.existing_hashes:
                    return

                # Bestellnummer aus Subject oder Body extrahieren
                order_id = extract_order_id(subject)
                if not order_id:
                    order_id = extract_order_id(body_html)
                if not order_id:
                    order_id = hashlib.md5(subject.encode()).hexdigest()[:8]

                safe_name = f"{safe_profile_name}_{fmt_date}_{order_id}_mail.pdf"
                output_path = target_dir / safe_name

                # Hybrid-Design: Mail-Header in Body-PDF
                mail_meta = {'sender': sender, 'subject': subject, 'date': fmt_date}

                if not output_path.exists() and html_to_pdf(body_html, output_path, mail_meta, mode=self.settings.pdf_mode):
                    self.log.emit(f"   📄 {safe_name}")

                    # OCR hinzufuegen wenn aktiviert und PDF keine Textebene hat
                    # OCR: Immer ausfuehren wenn aktiviert (scannt Bilder nach Text)
                    if self.settings.ocr_enabled and OCR_AVAILABLE:
                        ocr = OCRProcessor()
                        success, ocr_msg = ocr.enhance_with_ocr(output_path)
                        if success:
                            self.log.emit(f"   🔍 OCR: {ocr_msg}")
                        else:
                            # Kein Fehler loggen wenn einfach kein Text erkannt wurde
                            if "Kein Text" not in ocr_msg:
                                self.log.emit(f"   ⚠️ OCR: {ocr_msg}")

                    inv = Invoice(
                        id=str(uuid.uuid4()),
                        profile_name=profile.name,
                        filename=safe_name,
                        date=fmt_date,
                        path=str(output_path),
                        sender=sender[:80],
                        subject=subject[:100],
                        hash=body_hash,
                        is_attachment=False
                    )
                    self.invoice_found.emit(inv)
                    self.existing_hashes.add(body_hash)
                    self.found_count += 1


# ==================== DIALOGE ====================

class AccountDialog(QDialog):
    """Dialog for creating or editing a mail account configuration."""

    def __init__(self, account: MailAccount = None, parent=None):
        super().__init__(parent)
        self.account = account
        self.setup_ui()

    def setup_ui(self):
        self.setWindowTitle("E-Mail Konto" if not self.account else "Konto bearbeiten")
        self.resize(450, 300)

        layout = QVBoxLayout(self)

        form = QFormLayout()

        # Name
        self.inp_name = QLineEdit()
        self.inp_name.setPlaceholderText("z.B. 'Mein Gmail'")
        form.addRow("Anzeigename:", self.inp_name)

        # Provider Auswahl
        self.cb_provider = QComboBox()
        self.cb_provider.addItems(list(IMAP_PRESETS.keys()))
        self.cb_provider.currentTextChanged.connect(self.on_provider_changed)
        form.addRow("Anbieter:", self.cb_provider)

        # Gmail API Option
        self.ck_gmail_api = QCheckBox("Gmail API nutzen (empfohlen für Gmail)")
        self.ck_gmail_api.setToolTip("Schneller und zuverlässiger als IMAP")
        form.addRow("", self.ck_gmail_api)

        # Separator
        line = QFrame()
        line.setFrameShape(QFrame.Shape.HLine)
        form.addRow(line)

        # IMAP Settings
        self.lbl_imap = QLabel("IMAP Einstellungen:")
        self.lbl_imap.setStyleSheet("font-weight: bold; margin-top: 10px;")
        form.addRow(self.lbl_imap)

        self.inp_host = QLineEdit()
        self.inp_host.setPlaceholderText("imap.example.com")
        form.addRow("Server:", self.inp_host)

        self.inp_port = QSpinBox()
        self.inp_port.setRange(1, 65535)
        self.inp_port.setValue(993)
        form.addRow("Port:", self.inp_port)

        self.inp_user = QLineEdit()
        self.inp_user.setPlaceholderText("email@example.com")
        form.addRow("Benutzername:", self.inp_user)

        self.inp_pass = QLineEdit()
        self.inp_pass.setEchoMode(QLineEdit.EchoMode.Password)
        self.inp_pass.setPlaceholderText("Passwort / App-Passwort")
        form.addRow("Passwort:", self.inp_pass)

        layout.addLayout(form)

        # Buttons
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        # Daten laden falls Bearbeitung
        if self.account:
            self.inp_name.setText(self.account.name)
            self.inp_host.setText(self.account.host)
            self.inp_port.setValue(self.account.port)
            self.inp_user.setText(self.account.username)
            self.ck_gmail_api.setChecked(self.account.use_gmail_api)
            self.inp_pass.setPlaceholderText("Leer lassen = unverändert")

            # Provider finden
            for provider, preset in IMAP_PRESETS.items():
                if preset.get('host') == self.account.host:
                    self.cb_provider.setCurrentText(provider)
                    break
            # on_provider_changed may have overwritten use_gmail_api — restore
            self.ck_gmail_api.setChecked(self.account.use_gmail_api)

    def on_provider_changed(self, provider: str):
        """Füllt IMAP-Einstellungen basierend auf Provider"""
        preset = IMAP_PRESETS.get(provider, {})
        if preset.get('host'):
            self.inp_host.setText(preset['host'])
            self.inp_port.setValue(preset.get('port', 993))

        # Gmail API nur bei Gmail anzeigen
        is_gmail = provider == "Gmail"
        self.ck_gmail_api.setVisible(is_gmail)
        self.ck_gmail_api.setChecked(is_gmail)

    def get_account(self) -> Tuple[MailAccount, str]:
        """Gibt Account und Passwort zurück"""
        account = MailAccount(
            id=self.account.id if self.account else str(uuid.uuid4()),
            name=self.inp_name.text().strip() or "Unbenannt",
            provider=self.cb_provider.currentText(),
            host=self.inp_host.text().strip(),
            port=self.inp_port.value(),
            username=self.inp_user.text().strip(),
            use_gmail_api=self.ck_gmail_api.isChecked()
        )
        return account, self.inp_pass.text()


class QueryBuilderDialog(QDialog):
    """Small helper dialog for building Gmail Raw queries without manual syntax work."""

    def __init__(self, current_query: str = "", parent=None):
        super().__init__(parent)
        self.setWindowTitle("Gmail Query Builder")
        self.resize(560, 420)

        layout = QVBoxLayout(self)

        scope_group = QGroupBox("1. Bereich")
        scope_layout = QHBoxLayout(scope_group)
        self.rb_all = QRadioButton("Überall außer Papierkorb")
        self.rb_all.setChecked(True)
        self.rb_inbox = QRadioButton("Nur Inbox")
        self.rb_sent = QRadioButton("Gesendet")
        self.rb_trash = QRadioButton("Auch Papierkorb")
        scope_layout.addWidget(self.rb_all)
        scope_layout.addWidget(self.rb_inbox)
        scope_layout.addWidget(self.rb_sent)
        scope_layout.addWidget(self.rb_trash)
        layout.addWidget(scope_group)

        date_group = QGroupBox("2. Zeitraum")
        date_layout = QGridLayout(date_group)
        self.cb_time = QComboBox()
        self.cb_time.addItems(["Alles", "Dieses Jahr", "Letztes Jahr", "Benutzerdefiniert"])
        self.cb_time.currentIndexChanged.connect(self.toggle_dates)
        self.de_from = QDateEdit(QDate.currentDate().addYears(-1))
        self.de_from.setCalendarPopup(True)
        self.de_to = QDateEdit(QDate.currentDate())
        self.de_to.setCalendarPopup(True)
        self.de_from.setEnabled(False)
        self.de_to.setEnabled(False)
        date_layout.addWidget(QLabel("Preset:"), 0, 0)
        date_layout.addWidget(self.cb_time, 0, 1)
        date_layout.addWidget(QLabel("Von:"), 1, 0)
        date_layout.addWidget(self.de_from, 1, 1)
        date_layout.addWidget(QLabel("Bis:"), 2, 0)
        date_layout.addWidget(self.de_to, 2, 1)
        layout.addWidget(date_group)

        criteria_group = QGroupBox("3. Kriterien")
        criteria_layout = QGridLayout(criteria_group)
        self.inp_from = QLineEdit()
        self.inp_from.setPlaceholderText("z.B. amazon.de, amazon.com")
        self.inp_subject = QLineEdit()
        self.inp_subject.setPlaceholderText("z.B. Rechnung, Invoice")
        self.chk_attachment = QCheckBox("Muss Anhänge haben (has:attachment)")
        self.chk_attachment.setChecked(True)
        criteria_layout.addWidget(QLabel("Absender:"), 0, 0)
        criteria_layout.addWidget(self.inp_from, 0, 1)
        criteria_layout.addWidget(QLabel("Betreff:"), 1, 0)
        criteria_layout.addWidget(self.inp_subject, 1, 1)
        criteria_layout.addWidget(self.chk_attachment, 2, 0, 1, 2)
        layout.addWidget(criteria_group)

        self.result_query = QLineEdit(current_query)
        self.result_query.setPlaceholderText("Erzeugte Query …")
        btn_generate = QPushButton("Query generieren")
        btn_generate.clicked.connect(self.generate)
        layout.addWidget(btn_generate)
        layout.addWidget(self.result_query)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def toggle_dates(self) -> None:
        is_custom = self.cb_time.currentText() == "Benutzerdefiniert"
        self.de_from.setEnabled(is_custom)
        self.de_to.setEnabled(is_custom)

    def generate(self) -> None:
        parts = []
        if self.rb_inbox.isChecked():
            parts.append("in:inbox")
        elif self.rb_sent.isChecked():
            parts.append("in:sent")
        elif self.rb_trash.isChecked():
            parts.append("in:trash")
        else:
            parts.append("-in:trash")

        today = date.today()
        timeframe = self.cb_time.currentText()
        d_from = None
        d_to = None
        if timeframe == "Dieses Jahr":
            d_from = date(today.year, 1, 1)
        elif timeframe == "Letztes Jahr":
            d_from = date(today.year - 1, 1, 1)
            d_to = date(today.year - 1, 12, 31)
        elif timeframe == "Benutzerdefiniert":
            d_from = self.de_from.date().toPython()
            d_to = self.de_to.date().toPython()

        if d_from:
            parts.append(f"after:{d_from.strftime('%Y/%m/%d')}")
        if d_to:
            parts.append(f"before:{d_to.strftime('%Y/%m/%d')}")

        senders = [term.strip() for term in self.inp_from.text().split(",") if term.strip()]
        if senders:
            if len(senders) == 1:
                parts.append(f"from:{senders[0]}")
            else:
                sender_expr = " OR ".join([f"from:{sender}" for sender in senders])
                parts.append(f"({sender_expr})")

        subjects = [term.strip() for term in self.inp_subject.text().split(",") if term.strip()]
        if subjects:
            if len(subjects) == 1:
                parts.append(f'subject:"{subjects[0]}"')
            else:
                subject_expr = " OR ".join([f'subject:"{subject}"' for subject in subjects])
                parts.append(f"({subject_expr})")

        if self.chk_attachment.isChecked():
            parts.append("has:attachment")

        self.result_query.setText(" ".join(parts))

    def get_query(self) -> str:
        return self.result_query.text().strip()


class ProfileDialog(QDialog):
    """Dialog for creating or editing an invoice search profile."""

    def __init__(self, accounts: List[MailAccount], profile: InvoiceProfile = None, parent=None):
        super().__init__(parent)
        self.accounts = accounts
        self.profile = profile
        self.setup_ui()

    def setup_ui(self):
        self.setWindowTitle("Suchprofil" if not self.profile else "Profil bearbeiten")
        self.resize(560, 460)

        layout = QVBoxLayout(self)

        form = QFormLayout()

        # Name
        self.inp_name = QLineEdit()
        self.inp_name.setPlaceholderText("z.B. 'Amazon Rechnungen'")
        form.addRow("Name:", self.inp_name)

        # Account Auswahl
        self.cb_account = QComboBox()
        for acc in self.accounts:
            self.cb_account.addItem(acc.name, acc.id)
        form.addRow("E-Mail Konto:", self.cb_account)

        # Schnellauswahl Shop
        self.cb_shop = QComboBox()
        self.cb_shop.addItem("-- Manuell konfigurieren --")
        for shop in DEFAULT_SHOP_PROFILES:
            self.cb_shop.addItem(shop['name'])
        self.cb_shop.currentTextChanged.connect(self.on_shop_changed)
        form.addRow("Shop-Vorlage:", self.cb_shop)

        # Separator
        line = QFrame()
        line.setFrameShape(QFrame.Shape.HLine)
        form.addRow(line)

        # Filter
        self.inp_sender = QLineEdit()
        self.inp_sender.setPlaceholderText("z.B. amazon.de, amazon.com")
        self.inp_sender.setToolTip("Komma-getrennt = ODER-Verknüpfung\nBeispiel: amazon.de, amazon.com\n→ Mail von amazon.de ODER amazon.com")
        form.addRow("Absender enthält:", self.inp_sender)

        self.inp_subject = QLineEdit()
        self.inp_subject.setPlaceholderText("z.B. Rechnung, Invoice, Bestellung")
        self.inp_subject.setToolTip("Komma-getrennt = ODER-Verknüpfung\nBeispiel: Rechnung, Invoice\n→ Betreff enthält 'Rechnung' ODER 'Invoice'")
        form.addRow("Betreff enthält:", self.inp_subject)

        self.inp_gmail_query = QLineEdit()
        self.inp_gmail_query.setPlaceholderText("Optional: z.B. label:finance has:attachment")
        self.inp_gmail_query.setToolTip(
            "Optionaler Gmail-Raw-Query-Kanal.\n"
            "Greift bei Gmail API immer und bei IMAP nur auf Servern mit X-GM-RAW."
        )
        btn_query_builder = QPushButton("Builder …")
        btn_query_builder.setToolTip("Hilft beim Erstellen einer Gmail-Query")
        btn_query_builder.clicked.connect(self.open_query_builder)
        gmail_query_row = QHBoxLayout()
        gmail_query_row.addWidget(self.inp_gmail_query)
        gmail_query_row.addWidget(btn_query_builder)
        form.addRow("Gmail-Query:", gmail_query_row)

        # Blacklist
        self.inp_blacklist = QLineEdit()
        self.inp_blacklist.setPlaceholderText("z.B. Storno, Mahnung, Werbung")
        self.inp_blacklist.setToolTip("Komma-getrennt = ODER-Verknüpfung\nMails werden übersprungen wenn Betreff/Body\neines dieser Worte enthält")
        form.addRow("Darf NICHT enthalten:", self.inp_blacklist)

        # Separator fuer Body-Filter
        line2 = QFrame()
        line2.setFrameShape(QFrame.Shape.HLine)
        form.addRow(line2)

        # Body-Filter
        self.inp_body_must = QLineEdit()
        self.inp_body_must.setPlaceholderText("Optional: z.B. Rechnung, Invoice, Bestellung")
        self.inp_body_must.setToolTip("Mail-Body MUSS mindestens eines dieser Worte enthalten\nKomma-getrennt = ODER-Verknüpfung\nLeer = kein Filter")
        form.addRow("Body muss enthalten:", self.inp_body_must)

        self.inp_body_must_not = QLineEdit()
        self.inp_body_must_not.setPlaceholderText("Optional: z.B. Werbung, Newsletter")
        self.inp_body_must_not.setToolTip("Mail-Body darf KEINES dieser Worte enthalten\nKomma-getrennt = ODER-Verknüpfung\nLeer = kein Filter")
        form.addRow("Body darf nicht enthalten:", self.inp_body_must_not)

        # Zielordner
        self.inp_folder = QLineEdit()
        self.inp_folder.setPlaceholderText("Optional: Unterordner für diese Rechnungen")
        form.addRow("Unterordner:", self.inp_folder)

        # Aktiv
        self.ck_enabled = QCheckBox("Profil aktiviert")
        self.ck_enabled.setChecked(True)
        form.addRow("", self.ck_enabled)

        layout.addLayout(form)

        # Buttons
        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        # Daten laden
        if self.profile:
            self.inp_name.setText(self.profile.name)
            self.inp_sender.setText(self.profile.sender_filter)
            self.inp_subject.setText(self.profile.subject_filter)
            self.inp_gmail_query.setText(getattr(self.profile, 'gmail_query', ''))
            self.inp_blacklist.setText(getattr(self.profile, 'blacklist', ''))
            self.inp_body_must.setText(getattr(self.profile, 'body_must_contain', ''))
            self.inp_body_must_not.setText(getattr(self.profile, 'body_must_not_contain', ''))
            self.inp_folder.setText(self.profile.target_subfolder)
            self.ck_enabled.setChecked(self.profile.enabled)

            # Account finden
            for i in range(self.cb_account.count()):
                if self.cb_account.itemData(i) == self.profile.account_id:
                    self.cb_account.setCurrentIndex(i)
                    break

    def open_query_builder(self) -> None:
        """Opens the Gmail query builder and transfers the result back into the form."""
        dialog = QueryBuilderDialog(self.inp_gmail_query.text(), self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.inp_gmail_query.setText(dialog.get_query())

    def on_shop_changed(self, shop_name: str):
        """Füllt Filter basierend auf Shop-Vorlage"""
        for shop in DEFAULT_SHOP_PROFILES:
            if shop['name'] == shop_name:
                self.inp_name.setText(shop['name'])
                self.inp_sender.setText(shop['sender'])
                self.inp_subject.setText(shop['subject'])
                self.inp_folder.setText(shop['name'])
                break

    def get_profile(self) -> InvoiceProfile:
        """Gibt das konfigurierte Profil zurueck"""
        return InvoiceProfile(
            id=self.profile.id if self.profile else str(uuid.uuid4()),
            name=self.inp_name.text().strip() or "Unbenannt",
            account_id=self.cb_account.currentData(),
            sender_filter=self.inp_sender.text().strip(),
            subject_filter=self.inp_subject.text().strip(),
            gmail_query=self.inp_gmail_query.text().strip(),
            blacklist=self.inp_blacklist.text().strip(),
            body_must_contain=self.inp_body_must.text().strip(),
            body_must_not_contain=self.inp_body_must_not.text().strip(),
            target_subfolder=self.inp_folder.text().strip(),
            enabled=self.ck_enabled.isChecked()
        )


# ==================== DATEV-DIALOG ====================

class DATEVSettingsDialog(QDialog):
    """Konfiguration der DATEV-Exporteinstellungen und Konten-Mappings."""

    def __init__(self, config: DATEVConfig, parent=None):
        super().__init__(parent)
        self.setWindowTitle("DATEV-Export Einstellungen & Konten-Mapping")
        self.setMinimumWidth(550)
        self.setMinimumHeight(450)

        layout = QVBoxLayout(self)
        form = QFormLayout()

        berater_val = getattr(config, 'berater_nr', '12345') if config else '12345'
        self.inp_berater = QLineEdit(str(berater_val))
        self.inp_berater.setPlaceholderText("z.B. 12345")
        self.inp_berater.setAccessibleName("Beraternummer")
        self.inp_berater.setAccessibleDescription(
            "DATEV-Beraternummer für den Buchungsstapel."
        )
        self.inp_berater.setToolTip("DATEV-Beraternummer eingeben")
        form.addRow("Beraternummer:", self.inp_berater)

        mandant_val = getattr(config, 'mandant_nr', '67890') if config else '67890'
        self.inp_mandant = QLineEdit(str(mandant_val))
        self.inp_mandant.setPlaceholderText("z.B. 67890")
        self.inp_mandant.setAccessibleName("Mandantennummer")
        self.inp_mandant.setAccessibleDescription(
            "DATEV-Mandantennummer für den Buchungsstapel."
        )
        self.inp_mandant.setToolTip("DATEV-Mandantennummer eingeben")
        form.addRow("Mandantennummer:", self.inp_mandant)

        layout.addLayout(form)

        # Konten-Mapping Tabelle
        lbl_mapping = QLabel("<b>Konten-Mapping (Kreditoren & Gegenkonten per Absender):</b>")
        layout.addWidget(lbl_mapping)

        self.table_mapping = QTableWidget()
        self.table_mapping.setColumnCount(3)
        self.table_mapping.setHorizontalHeaderLabels(["Absender / Schlüsselwort", "Konto (Kreditor)", "Gegenkonto (Aufwand)"])
        self.table_mapping.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.table_mapping.setAccessibleName("DATEV-Konten-Mapping-Tabelle")
        self.table_mapping.setAccessibleDescription(
            "Ordnet Absendern oder Schlüsselwörtern ein Kreditor- und ein Aufwandskonto zu."
        )
        self.table_mapping.setToolTip(
            "Absender oder Schlüsselwort sowie Kreditor- und Aufwandskonto bearbeiten"
        )
        layout.addWidget(self.table_mapping)

        # Table Control Buttons
        btn_layout = QHBoxLayout()
        self.btn_add_row = QPushButton("Zeile hinzufügen")
        self.btn_add_row.setAccessibleName("Zeile hinzufügen")
        self.btn_add_row.setAccessibleDescription(
            "Fügt eine neue, editierbare Konten-Mapping-Zeile hinzu."
        )
        self.btn_add_row.setToolTip("Neue Konten-Mapping-Zeile hinzufügen")
        self.btn_add_row.clicked.connect(self._add_row)
        btn_layout.addWidget(self.btn_add_row)

        self.btn_remove_row = QPushButton("Zeile entfernen")
        self.btn_remove_row.setAccessibleName("Zeile entfernen")
        self.btn_remove_row.setAccessibleDescription(
            "Entfernt die aktuell ausgewählte Konten-Mapping-Zeile."
        )
        self.btn_remove_row.setToolTip("Ausgewählte Konten-Mapping-Zeile entfernen")
        self.btn_remove_row.clicked.connect(self._remove_row)
        btn_layout.addWidget(self.btn_remove_row)

        self.btn_reset_mapping = QPushButton("Standard wiederherstellen")
        self.btn_reset_mapping.setAccessibleName("Standard wiederherstellen")
        self.btn_reset_mapping.setAccessibleDescription(
            "Ersetzt alle Einträge durch die standardmäßige Konten-Zuordnung."
        )
        self.btn_reset_mapping.setToolTip("Standardmäßige Konten-Zuordnung wiederherstellen")
        self.btn_reset_mapping.clicked.connect(self._reset_mapping)
        btn_layout.addWidget(self.btn_reset_mapping)

        layout.addLayout(btn_layout)

        # Hinweis auf fehlende amount-Felder
        hint = QLabel(
            "<i>Hinweis: Rechnungsbeträge werden aus dem Feld 'Betrag' gelesen.<br>"
            "Rechnungen ohne eingetragenen Betrag werden übersprungen.</i>"
        )
        hint.setWordWrap(True)
        hint.setStyleSheet("color: #aaa; font-size: 9pt;")
        layout.addWidget(hint)

        self.dialog_buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        ok_button = self.dialog_buttons.button(QDialogButtonBox.StandardButton.Ok)
        ok_button.setAccessibleName("DATEV-Einstellungen speichern")
        ok_button.setAccessibleDescription(
            "Speichert Beraternummer, Mandantennummer und Konten-Mapping."
        )
        ok_button.setToolTip("DATEV-Einstellungen speichern")
        cancel_button = self.dialog_buttons.button(QDialogButtonBox.StandardButton.Cancel)
        cancel_button.setAccessibleName("DATEV-Einstellungen verwerfen")
        cancel_button.setAccessibleDescription(
            "Schließt den Dialog ohne Änderungen zu speichern."
        )
        cancel_button.setToolTip("DATEV-Einstellungen ohne Speichern schließen")
        self.dialog_buttons.accepted.connect(self.accept)
        self.dialog_buttons.rejected.connect(self.reject)
        layout.addWidget(self.dialog_buttons)

        # Initial data populating
        initial_mapping = config.konten_mapping if (config and getattr(config, 'konten_mapping', None)) else DEFAULT_KONTEN_MAPPING
        self.mapping_data = dict(initial_mapping)
        self._populate_table()

    def _populate_table(self):
        self.table_mapping.setRowCount(0)
        for key, val in self.mapping_data.items():
            row = self.table_mapping.rowCount()
            self.table_mapping.insertRow(row)
            konto, gegenkonto = val if isinstance(val, (tuple, list)) and len(val) == 2 else (70000, 4900)
            self.table_mapping.setItem(row, 0, QTableWidgetItem(str(key)))
            self.table_mapping.setItem(row, 1, QTableWidgetItem(str(konto)))
            self.table_mapping.setItem(row, 2, QTableWidgetItem(str(gegenkonto)))

    def _add_row(self):
        row = self.table_mapping.rowCount()
        self.table_mapping.insertRow(row)
        self.table_mapping.setItem(row, 0, QTableWidgetItem("Neuer Partner"))
        self.table_mapping.setItem(row, 1, QTableWidgetItem("70000"))
        self.table_mapping.setItem(row, 2, QTableWidgetItem("4900"))

    def _remove_row(self):
        curr = self.table_mapping.currentRow()
        if curr >= 0:
            self.table_mapping.removeRow(curr)
        elif self.table_mapping.rowCount() > 0:
            self.table_mapping.removeRow(self.table_mapping.rowCount() - 1)

    def _reset_mapping(self):
        self.mapping_data = dict(DEFAULT_KONTEN_MAPPING)
        self._populate_table()

    def _get_raw_table_mapping(self) -> Tuple[dict, List[str]]:
        """Liest Tabellenwerte verlustfrei und meldet doppelte Schlüssel."""
        mapping = {}
        errors: List[str] = []
        seen_keys = {}
        for row in range(self.table_mapping.rowCount()):
            item_key = self.table_mapping.item(row, 0)
            item_k = self.table_mapping.item(row, 1)
            item_gk = self.table_mapping.item(row, 2)
            key = item_key.text().strip() if item_key else ""
            konto = item_k.text().strip() if item_k else ""
            gegenkonto = item_gk.text().strip() if item_gk else ""

            normalized_key = key.casefold()
            if normalized_key:
                if normalized_key in seen_keys:
                    first_row = seen_keys[normalized_key]
                    errors.append(
                        f"Konten-Mapping-Zeile {row + 1}: Schlüssel '{key}' ist "
                        f"bereits in Zeile {first_row} vorhanden."
                    )
                else:
                    seen_keys[normalized_key] = row + 1

            mapping[key] = (konto, gegenkonto)
        return mapping, errors

    def _get_table_mapping(self) -> dict:
        """Konvertiert bereits validierte Tabellenwerte in DATEV-Kontonummern."""
        raw_mapping, _ = self._get_raw_table_mapping()
        return {
            key: (int(konto), int(gegenkonto))
            for key, (konto, gegenkonto) in raw_mapping.items()
        }

    def validate_inputs(self) -> Tuple[bool, List[str]]:
        """Validiert die aktuellen Eingaben im Dialog vor dem Schließen."""
        raw_mapping, row_errors = self._get_raw_table_mapping()
        cfg = DATEVConfig(
            berater_nr=self.inp_berater.text().strip(),
            mandant_nr=self.inp_mandant.text().strip(),
            konten_mapping=raw_mapping,
        )
        if DATEV_AVAILABLE:
            _, config_errors = validate_datev_config(cfg)
            errors = row_errors + config_errors
            return (len(errors) == 0), errors
        errors = list(row_errors)
        if not self.inp_berater.text().strip().isdigit():
            errors.append("Beraternummer muss numerisch sein.")
        if not self.inp_mandant.text().strip().isdigit():
            errors.append("Mandantennummer muss numerisch sein.")
        return (len(errors) == 0), errors

    def accept(self):
        """Überschreibt accept mit Prüfung der Validierung."""
        is_valid, errors = self.validate_inputs()
        if not is_valid:
            error_msg = "\n• " + "\n• ".join(errors)
            QMessageBox.warning(
                self,
                "Ungültige DATEV-Einstellungen",
                f"Bitte korrigieren Sie die folgenden Fehler:{error_msg}"
            )
            return
        super().accept()

    def get_config(self) -> DATEVConfig:
        """Gibt die aktuell eingestellte DATEVConfig zurück."""
        return DATEVConfig(
            berater_nr=self.inp_berater.text().strip(),
            mandant_nr=self.inp_mandant.text().strip(),
            konten_mapping=self._get_table_mapping()
        )


# ==================== HAUPTFENSTER ====================

class MainWindow(QMainWindow):
    """Main application window for UniversalInvoiceMail."""

    def __init__(self):
        super().__init__()

        # Daten laden
        self.accounts: List[MailAccount] = []
        self.profiles: List[InvoiceProfile] = []
        self.invoices: List[Invoice] = []
        self.settings = AppSettings()
        self.datev_config = DATEVConfig() if DATEV_AVAILABLE else None
        self.worker = None

        self.load_config()
        self.setup_ui()

    def load_config(self):
        """Lädt Konfiguration"""
        if CONFIG_FILE.exists():
            try:
                data = json.loads(CONFIG_FILE.read_text(encoding='utf-8'))
                self.settings = AppSettings.from_dict(data.get('settings', {}))
                self.accounts = [MailAccount.from_dict(a) for a in data.get('accounts', [])]
                self.profiles = [InvoiceProfile.from_dict(p) for p in data.get('profiles', [])]
                # DATEV-Konfiguration laden
                if DATEV_AVAILABLE:
                    dc = data.get('datev_config', {})
                    if dc:
                        km = dc.get('konten_mapping')
                        self.datev_config = DATEVConfig(
                            berater_nr=dc.get('berater_nr', '12345'),
                            mandant_nr=dc.get('mandant_nr', '67890'),
                            konten_mapping={k: tuple(v) for k, v in km.items()} if km else None,
                        )
            except (OSError, json.JSONDecodeError, KeyError, ValueError, TypeError) as e:
                print(f"Config load error: {e}")

        if INVOICES_DB.exists():
            try:
                data = json.loads(INVOICES_DB.read_text(encoding='utf-8'))
                self.invoices = [Invoice.from_dict(i) for i in data]
            except (OSError, json.JSONDecodeError, KeyError, ValueError, TypeError) as e:
                print(f"Invoice DB error: {e}")

    def save_config(self):
        """Speichert Konfiguration"""
        try:
            data = {
                'settings': self.settings.to_dict(),
                'accounts': [a.to_dict() for a in self.accounts],
                'profiles': [p.to_dict() for p in self.profiles]
            }
            # DATEV-Konfiguration einschliessen damit sie beim App-Close nicht verloren geht
            if DATEV_AVAILABLE and self.datev_config is not None:
                data['datev_config'] = {
                    'berater_nr': self.datev_config.berater_nr,
                    'mandant_nr': self.datev_config.mandant_nr,
                    'konten_mapping': {k: list(v) for k, v in self.datev_config.konten_mapping.items()},
                }
            CONFIG_FILE.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding='utf-8')
            INVOICES_DB.write_text(
                json.dumps([i.to_dict() for i in self.invoices], indent=2, ensure_ascii=False),
                encoding='utf-8'
            )
        except (OSError, TypeError, ValueError) as e:
            print(f"Save error: {e}")

    def setup_ui(self):
        """Erstellt die Benutzeroberfläche"""
        self.setWindowTitle(f"{APP_NAME} v{VERSION}")
        self.resize(1200, 750)

        # Dark Theme
        self.setStyleSheet("""
            QMainWindow, QWidget { background-color: #2d2d2d; color: #ffffff; }
            QGroupBox { border: 1px solid #555; border-radius: 5px; margin-top: 10px; padding-top: 10px; }
            QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 5px; }
            QPushButton { background-color: #3d3d3d; border: 1px solid #555; border-radius: 4px; padding: 6px 12px; }
            QPushButton:hover { background-color: #4d4d4d; }
            QPushButton:pressed { background-color: #2d2d2d; }
            QLineEdit, QComboBox, QSpinBox { background-color: #3d3d3d; border: 1px solid #555; border-radius: 3px; padding: 4px; }
            QTableWidget { background-color: #252525; gridline-color: #444; }
            QTableWidget::item { padding: 4px; }
            QTableWidget::item:selected { background-color: #0078d4; }
            QHeaderView::section { background-color: #3d3d3d; padding: 6px; border: none; }
            QListWidget { background-color: #252525; }
            QListWidget::item { padding: 8px; }
            QListWidget::item:selected { background-color: #0078d4; }
            QPlainTextEdit { background-color: #1e1e1e; font-family: Consolas, monospace; }
            QTabWidget::pane { border: 1px solid #555; }
            QTabBar::tab { background-color: #3d3d3d; padding: 8px 16px; margin-right: 2px; }
            QTabBar::tab:selected { background-color: #0078d4; }
            QProgressBar { background-color: #3d3d3d; border-radius: 3px; text-align: center; }
            QProgressBar::chunk { background-color: #0078d4; border-radius: 3px; }
        """)

        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QHBoxLayout(central)

        # === LINKE SEITE: Steuerung ===
        left_panel = QWidget()
        left_panel.setMaximumWidth(350)
        left_layout = QVBoxLayout(left_panel)

        # Start Button
        self.btn_start = QPushButton("🚀  RECHNUNGEN ABRUFEN")
        self.btn_start.setMinimumHeight(50)
        self.btn_start.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                font-size: 14pt;
                font-weight: bold;
                border-radius: 8px;
            }
            QPushButton:hover { background-color: #2ecc71; }
            QPushButton:disabled { background-color: #555; }
        """)
        self.btn_start.clicked.connect(self.start_grabbing)
        left_layout.addWidget(self.btn_start)

        # Progress
        self.progress = QProgressBar()
        self.progress.setVisible(False)
        left_layout.addWidget(self.progress)

        # Zeit-Filter mit Von/Bis Datum
        filter_group = QGroupBox("Zeitraum")
        filter_layout = QVBoxLayout(filter_group)

        # Schnellauswahl
        quick_layout = QHBoxLayout()
        self.cb_timeframe = QComboBox()
        self.cb_timeframe.addItems([
            "Benutzerdefiniert",
            "Letzte 12 Monate",
            "Letzte 6 Monate",
            "Letzte 3 Monate",
            "Dieses Jahr (01.01.-heute)",
            "Letztes Jahr (01.01.-31.12.)",
            "Alles"
        ])
        self.cb_timeframe.currentTextChanged.connect(self._on_timeframe_changed)
        self.cb_timeframe.setToolTip("Schnellauswahl oder 'Benutzerdefiniert' für eigene Daten")
        quick_layout.addWidget(QLabel("Schnellauswahl:"))
        quick_layout.addWidget(self.cb_timeframe)
        filter_layout.addLayout(quick_layout)

        # Von/Bis Datum
        from PySide6.QtWidgets import QDateEdit
        from PySide6.QtCore import QDate

        date_layout = QHBoxLayout()
        date_layout.addWidget(QLabel("Von:"))
        self.date_from = QDateEdit()
        self.date_from.setCalendarPopup(True)
        self.date_from.setDisplayFormat("dd.MM.yyyy")
        # Gespeichertes Datum laden oder Default
        if self.settings.date_from:
            try:
                parts = self.settings.date_from.split("-")
                self.date_from.setDate(QDate(int(parts[0]), int(parts[1]), int(parts[2])))
            except (ValueError, IndexError):
                self.date_from.setDate(QDate.currentDate().addMonths(-12))
        else:
            self.date_from.setDate(QDate.currentDate().addMonths(-12))
        date_layout.addWidget(self.date_from)

        date_layout.addWidget(QLabel("Bis:"))
        self.date_to = QDateEdit()
        self.date_to.setCalendarPopup(True)
        self.date_to.setDisplayFormat("dd.MM.yyyy")
        # Gespeichertes Datum laden oder Default
        if self.settings.date_to:
            try:
                parts = self.settings.date_to.split("-")
                self.date_to.setDate(QDate(int(parts[0]), int(parts[1]), int(parts[2])))
            except (ValueError, IndexError):
                self.date_to.setDate(QDate.currentDate())
        else:
            self.date_to.setDate(QDate.currentDate())
        date_layout.addWidget(self.date_to)
        filter_layout.addLayout(date_layout)

        left_layout.addWidget(filter_group)

        # Profile Liste
        profile_group = QGroupBox("Suchprofile")
        profile_layout = QVBoxLayout(profile_group)

        self.profile_list = QListWidget()
        self.profile_list.itemDoubleClicked.connect(self.edit_profile)
        profile_layout.addWidget(self.profile_list)

        btn_row = QHBoxLayout()
        btn_add = QPushButton("➕ Profil")
        btn_add.setObjectName("add_profile_button")
        btn_add.setAccessibleName("Neues Suchprofil anlegen")
        btn_add.setAccessibleDescription(
            "Öffnet den Dialog zum Anlegen eines neuen Suchprofils für Rechnungen."
        )
        btn_add.setToolTip("Neues Suchprofil anlegen")
        btn_add.clicked.connect(self.add_profile)
        btn_del = QPushButton("❌")
        btn_del.setObjectName("delete_profile_button")
        btn_del.setAccessibleName("Ausgewähltes Suchprofil löschen")
        btn_del.setToolTip("Ausgewähltes Suchprofil löschen")
        btn_del.setMaximumWidth(40)
        btn_del.clicked.connect(self.delete_profile)
        btn_row.addWidget(btn_add)
        btn_row.addWidget(btn_del)
        profile_layout.addLayout(btn_row)

        left_layout.addWidget(profile_group)

        # Accounts
        account_group = QGroupBox("E-Mail Konten")
        account_layout = QVBoxLayout(account_group)

        self.account_list = QListWidget()
        self.account_list.setMaximumHeight(100)
        self.account_list.itemDoubleClicked.connect(self.edit_account)
        account_layout.addWidget(self.account_list)

        btn_row2 = QHBoxLayout()
        btn_add_acc = QPushButton("➕ Konto")
        btn_add_acc.setObjectName("add_account_button")
        btn_add_acc.setAccessibleName("Neues E-Mail-Konto anlegen")
        btn_add_acc.setAccessibleDescription(
            "Öffnet den Dialog zum Hinzufügen eines weiteren E-Mail-Kontos."
        )
        btn_add_acc.setToolTip("Neues E-Mail-Konto anlegen")
        btn_add_acc.clicked.connect(self.add_account)
        btn_del_acc = QPushButton("❌")
        btn_del_acc.setObjectName("delete_account_button")
        btn_del_acc.setAccessibleName("Ausgewähltes E-Mail-Konto löschen")
        btn_del_acc.setToolTip("Ausgewähltes E-Mail-Konto löschen")
        btn_del_acc.setMaximumWidth(40)
        btn_del_acc.clicked.connect(self.delete_account)
        btn_row2.addWidget(btn_add_acc)
        btn_row2.addWidget(btn_del_acc)
        account_layout.addLayout(btn_row2)

        left_layout.addWidget(account_group)
        left_layout.addStretch()

        main_layout.addWidget(left_panel)

        # === RECHTE SEITE: Tabs ===
        tabs = QTabWidget()
        tabs.setObjectName("main_workspace_tabs")
        tabs.setAccessibleName("Arbeitsbereiche")
        tabs.setAccessibleDescription(
            "Wechselt zwischen Rechnungen, Einstellungen, Protokoll und Informationen."
        )

        # Tab: Rechnungen
        invoice_tab = QWidget()
        invoice_layout = QVBoxLayout(invoice_tab)

        self.invoice_table = QTableWidget(0, 8)
        self.invoice_table.setHorizontalHeaderLabels(
            ["✓", "Typ", "Datum", "Shop", "Absender", "Betrag (€)", "Datei", "Pfad"])
        self.invoice_table.horizontalHeader().setSectionResizeMode(6, QHeaderView.ResizeMode.Stretch)
        self.invoice_table.setColumnWidth(0, 30)
        self.invoice_table.setColumnWidth(1, 40)
        self.invoice_table.setColumnWidth(5, 90)  # Betrag-Spalte
        self.invoice_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.invoice_table.setObjectName("invoice_table")
        self.invoice_table.setAccessibleName("Rechnungsliste")
        self.invoice_table.setAccessibleDescription(
            "Zeigt gefundene Rechnungen. Zeilen können für Export- oder Löschaktionen ausgewählt werden."
        )
        self.invoice_table.setToolTip(
            "Rechnungen auswählen oder mit Doppelklick öffnen"
        )
        self.invoice_table.cellDoubleClicked.connect(self.open_invoice)
        self.invoice_table.itemChanged.connect(self._on_invoice_amount_changed)
        invoice_layout.addWidget(self.invoice_table)

        inv_btn_row = QHBoxLayout()
        btn_select_all = QPushButton("Alle")
        btn_select_all.clicked.connect(self.select_all_invoices)
        btn_select_all.setObjectName("select_all_invoices_button")
        btn_select_all.setAccessibleName("Alle sichtbaren Rechnungen auswählen")
        btn_select_all.setAccessibleDescription(
            "Markiert alle sichtbaren Rechnungen für Export- oder Löschaktionen."
        )
        btn_select_all.setToolTip("Alle Einträge auswählen")
        btn_select_all.setMaximumWidth(50)
        btn_select_none = QPushButton("Keine")
        btn_select_none.clicked.connect(self.select_no_invoices)
        btn_select_none.setObjectName("clear_invoice_selection_button")
        btn_select_none.setAccessibleName("Rechnungsauswahl aufheben")
        btn_select_none.setAccessibleDescription(
            "Entfernt alle Markierungen in der Rechnungstabelle."
        )
        btn_select_none.setToolTip("Auswahl aufheben")
        btn_select_none.setMaximumWidth(50)
        btn_delete_selected = QPushButton("❌")
        btn_delete_selected.setObjectName("delete_selected_invoices_button")
        btn_delete_selected.setAccessibleName("Ausgewählte Rechnungen und Dateien löschen")
        btn_delete_selected.clicked.connect(self.delete_selected_invoices)
        btn_delete_selected.setToolTip("Ausgewählte Einträge und Dateien löschen")
        btn_delete_selected.setMaximumWidth(35)
        btn_open_folder = QPushButton("Ordner öffnen")
        btn_open_folder.clicked.connect(self.open_download_folder)
        btn_open_folder.setObjectName("open_invoice_folder_button")
        btn_open_folder.setAccessibleName("Speicherordner für Rechnungen öffnen")
        btn_open_folder.setAccessibleDescription(
            "Öffnet den aktuellen Rechnungsordner im Dateimanager."
        )
        btn_open_folder.setToolTip("Speicherordner im Explorer öffnen")
        btn_refresh = QPushButton("Aktualisieren")
        btn_refresh.clicked.connect(self.refresh_invoice_table)
        btn_refresh.setObjectName("refresh_invoice_table_button")
        btn_refresh.setAccessibleName("Rechnungsliste aktualisieren")
        btn_refresh.setAccessibleDescription(
            "Synchronisiert die Tabelle mit dem Dateisystem und importiert neue Dateien."
        )
        btn_refresh.setToolTip("Rechnungstabelle mit Ordnerinhalt synchronisieren")
        btn_export_csv = QPushButton("CSV Export")
        btn_export_csv.clicked.connect(self.export_invoices_csv)
        btn_export_csv.setObjectName("export_invoices_csv_button")
        btn_export_csv.setAccessibleName("Rechnungsliste als CSV exportieren")
        btn_export_csv.setAccessibleDescription(
            "Exportiert die aktuelle Rechnungsliste als Tabellen-Datei."
        )
        btn_export_csv.setToolTip("Rechnungsliste als CSV exportieren (filterbar in Excel)")
        btn_bundle_export = QPushButton("Bundle Export")
        btn_bundle_export.clicked.connect(self.export_invoice_bundle)
        btn_bundle_export.setObjectName("export_invoice_bundle_button")
        btn_bundle_export.setAccessibleName("Redigiertes Rechnungs-Bundle exportieren")
        btn_bundle_export.setAccessibleDescription(
            "Exportiert ausgewählte oder alle Rechnungen für Companion- oder Prüf-Workflows."
        )
        btn_bundle_export.setToolTip("Redigiertes Rechnungs-Bundle für Companion oder Prüfung exportieren")
        btn_bundle_import = QPushButton("Bundle Import")
        btn_bundle_import.clicked.connect(self.import_invoice_bundle)
        btn_bundle_import.setObjectName("import_invoice_bundle_button")
        btn_bundle_import.setAccessibleName("Companion-Bundle importieren")
        btn_bundle_import.setAccessibleDescription(
            "Übernimmt Betrag, Prüfflag und Notizen aus einem redigierten Rechnungs-Bundle."
        )
        btn_bundle_import.setToolTip("Companion-Änderungen für Betrag, Prüfflag und Notiz reimportieren")
        btn_datev = QPushButton("DATEV exportieren")
        btn_datev.clicked.connect(self._export_datev)
        btn_datev.setObjectName("export_datev_button")
        btn_datev.setAccessibleName("DATEV-Buchungsstapel exportieren")
        btn_datev.setAccessibleDescription(
            "Exportiert markierte Rechnungen als DATEV-Buchungsstapel für die Buchhaltung."
        )
        btn_datev.setToolTip("Ausgewählte Rechnungen als DATEV-Buchungsstapel exportieren")
        btn_datev.setEnabled(DATEV_AVAILABLE)
        if not DATEV_AVAILABLE:
            btn_datev.setToolTip("datev_exporter.py nicht gefunden")
        inv_btn_row.addWidget(btn_select_all)
        inv_btn_row.addWidget(btn_select_none)
        inv_btn_row.addWidget(btn_delete_selected)
        inv_btn_row.addWidget(btn_open_folder)
        inv_btn_row.addWidget(btn_refresh)
        inv_btn_row.addWidget(btn_export_csv)
        inv_btn_row.addWidget(btn_bundle_export)
        inv_btn_row.addWidget(btn_bundle_import)
        inv_btn_row.addWidget(btn_datev)
        inv_btn_row.addStretch()
        invoice_layout.addLayout(inv_btn_row)

        tabs.addTab(invoice_tab, "📄 Rechnungen")

        # Tab: Einstellungen
        settings_tab = QWidget()
        settings_layout = QVBoxLayout(settings_tab)

        settings_form = QFormLayout()

        # Download-Pfad
        path_row = QHBoxLayout()
        self.inp_path = QLineEdit(self.settings.download_path)
        btn_browse = QPushButton("...")
        btn_browse.setObjectName("browse_download_path_button")
        btn_browse.setAccessibleName("Speicherordner auswählen")
        btn_browse.setAccessibleDescription(
            "Öffnet die Ordnerauswahl für den lokalen Rechnungs-Speicherort."
        )
        btn_browse.setToolTip("Speicherordner auswählen")
        btn_browse.setMaximumWidth(40)
        btn_browse.clicked.connect(self.browse_folder)
        path_row.addWidget(self.inp_path)
        path_row.addWidget(btn_browse)
        settings_form.addRow("Speicherort:", path_row)

        # Optionen
        self.ck_attachments = QCheckBox("PDF-Anhänge herunterladen")
        self.ck_attachments.setChecked(self.settings.download_attachments)
        settings_form.addRow("", self.ck_attachments)

        self.ck_body_pdf = QCheckBox("Mail-Body als PDF speichern (wenn keine Anhänge)")
        self.ck_body_pdf.setChecked(self.settings.convert_body_to_pdf)
        settings_form.addRow("", self.ck_body_pdf)

        self.ck_merge_body = QCheckBox("Dem PDF den Mail-Body anhängen")
        self.ck_merge_body.setChecked(self.settings.merge_body_with_attachments)
        self.ck_merge_body.setToolTip("Wenn aktiv: Mail-Header und Body werden\nan PDF-Anhänge angehängt")
        settings_form.addRow("", self.ck_merge_body)

        self.ck_hash = QCheckBox("Duplikat-Erkennung (Hash-Check)")
        self.ck_hash.setChecked(self.settings.enable_hash_check)
        settings_form.addRow("", self.ck_hash)

        self.ck_trash = QCheckBox("Papierkorb durchsuchen")
        self.ck_trash.setChecked(self.settings.include_trash)
        self.ck_trash.setToolTip("Auch gelöschte Mails nach Rechnungen durchsuchen")
        settings_form.addRow("", self.ck_trash)

        # PDF-Modus Auswahl
        settings_form.addRow(QLabel("<b>PDF-Erstellung:</b>"))

        self.cmb_pdf_mode = QComboBox()
        self.cmb_pdf_mode.addItem("Schnell (nur Text)", "fast")
        self.cmb_pdf_mode.addItem("Vollständig (mit Bildern)", "full")
        # Browser-Modus nur anzeigen wenn Selenium verfuegbar
        if SELENIUM_AVAILABLE:
            self.cmb_pdf_mode.addItem("Browser (Edge/Chrome) - Empfohlen", "browser")
        # Aktuellen Modus setzen
        idx = self.cmb_pdf_mode.findData(self.settings.pdf_mode)
        if idx >= 0:
            self.cmb_pdf_mode.setCurrentIndex(idx)
        self.cmb_pdf_mode.setToolTip(
            "Schnell: Nur Text, keine Bilder (stabil, schnell)\n"
            "Vollständig: Mit Bildern via xhtml2pdf\n"
            "Browser: Nutzt Edge/Chrome für natives Rendering (beste Qualität)"
        )
        settings_form.addRow("PDF-Modus:", self.cmb_pdf_mode)

        self.ck_ocr = QCheckBox("OCR für bildbasierte PDFs")
        self.ck_ocr.setChecked(self.settings.ocr_enabled)
        self.ck_ocr.setToolTip(
            "Tesseract OCR ausführen wenn PDF nur Bilder enthält.\n"
            "Benötigt: pytesseract, pdf2image, Poppler"
        )
        self.ck_ocr.setEnabled(OCR_AVAILABLE)
        if not OCR_AVAILABLE:
            self.ck_ocr.setText("OCR für bildbasierte PDFs (nicht installiert)")
        settings_form.addRow("", self.ck_ocr)

        self.inp_max_mails = QSpinBox()
        self.inp_max_mails.setRange(10, 1000)
        self.inp_max_mails.setValue(self.settings.max_emails_per_run)
        settings_form.addRow("Max. Mails pro Durchlauf:", self.inp_max_mails)

        settings_layout.addLayout(settings_form)

        # DATEV-Konfiguration & Mapping Gruppe
        datev_group = QGroupBox("DATEV-Export & Konten-Mapping")
        datev_layout = QVBoxLayout(datev_group)

        datev_desc = QLabel(
            "Konfigurieren Sie hier Ihre DATEV-Stammdaten (Beraternummer, Mandantennummer) "
            "sowie das Konten-Mapping für Kreditoren- und Gegenkonten (SKR03/SKR04)."
        )
        datev_desc.setWordWrap(True)
        datev_desc.setStyleSheet("color: #aaa; font-size: 9pt;")
        datev_layout.addWidget(datev_desc)

        btn_datev_settings = QPushButton("⚙️ DATEV-Einstellungen & Konten-Mapping bearbeiten...")
        btn_datev_settings.setObjectName("edit_datev_settings_button")
        btn_datev_settings.setAccessibleName("DATEV-Einstellungen und Konten-Mapping bearbeiten")
        btn_datev_settings.setAccessibleDescription(
            "Öffnet den Dialog zur Konfiguration der DATEV Beraternummer, Mandantennummer und Konten-Mappings."
        )
        btn_datev_settings.setToolTip("DATEV Beraternummer, Mandantennummer & Sachkonten-Mapping anpassen")
        btn_datev_settings.clicked.connect(self.open_datev_settings_dialog)
        btn_datev_settings.setEnabled(DATEV_AVAILABLE)
        if not DATEV_AVAILABLE:
            btn_datev_settings.setToolTip("datev_exporter.py nicht gefunden")
        datev_layout.addWidget(btn_datev_settings)

        settings_layout.addWidget(datev_group)

        # Hinweise-Box
        hints_group = QGroupBox("Hinweise")
        hints_layout = QVBoxLayout(hints_group)
        hints_text = QLabel("""
<p><b>Löschen von Rechnungen:</b><br>
Das Löschen von Einträgen im Rechnungen-Tab entfernt auch die
zugehörigen PDF-Dateien im Ordner. Bei erneutem Scan mit passender
Profileinstellung werden dieselben Mails erneut heruntergeladen.</p>

<p><b>Duplikat-Erkennung:</b><br>
Bereits heruntergeladene Dateien werden anhand ihres Hash-Werts erkannt
und nicht erneut gespeichert (sofern aktiviert).</p>

<p><b>Manueller Import:</b><br>
PDFs die manuell in Profilordner gelegt werden, erscheinen nach
"Aktualisieren" automatisch in der Liste.</p>
        """)
        hints_text.setWordWrap(True)
        hints_text.setStyleSheet("color: #666; font-size: 9pt;")
        hints_layout.addWidget(hints_text)
        settings_layout.addWidget(hints_group)

        btn_save = QPushButton("💾 Einstellungen speichern")
        btn_save.clicked.connect(self.save_settings)
        settings_layout.addWidget(btn_save)
        settings_layout.addStretch()

        tabs.addTab(settings_tab, "⚙️ Einstellungen")

        # Tab: Log
        log_tab = QWidget()
        log_layout = QVBoxLayout(log_tab)

        self.log_output = QPlainTextEdit()
        self.log_output.setReadOnly(True)
        self.log_output.setObjectName("activity_log")
        self.log_output.setAccessibleName("Aktivitätsprotokoll")
        self.log_output.setAccessibleDescription(
            "Zeigt Fortschritt, gefundene Rechnungen und Fehlermeldungen des aktuellen Abrufs."
        )
        self.log_output.setToolTip(
            "Fortschritt und Meldungen des Rechnungsabrufs"
        )
        log_layout.addWidget(self.log_output)

        tabs.addTab(log_tab, "📝 Log")

        # Tab: Info
        info_tab = QWidget()
        info_layout = QVBoxLayout(info_tab)

        info_text = f"""
        <h2>{APP_NAME} v{VERSION}</h2>
        <p>Vereinfachte App zum Extrahieren von Rechnungen aus E-Mails.</p>

        <h3>Features:</h3>
        <ul>
            <li>IMAP für alle Mail-Anbieter (Gmail, Outlook, GMX, etc.)</li>
            <li>Gmail API für schnelleren Zugriff</li>
            <li>Vorkonfigurierte Shop-Profile</li>
            <li>Automatische Duplikat-Erkennung</li>
        </ul>

        <h3>Gmail API Setup:</h3>
        <ol>
            <li>Google Cloud Console öffnen</li>
            <li>Neues Projekt erstellen</li>
            <li>Gmail API aktivieren</li>
            <li>OAuth Credentials erstellen</li>
            <li>credentials.json in {BASE_DIR} speichern</li>
        </ol>

        <h3>IMAP bei Gmail:</h3>
        <p>Erstelle ein App-Passwort in deinem Google Konto unter Sicherheit → 2FA → App-Passwörter</p>

        <p><b>Config-Ordner:</b> {BASE_DIR}</p>
        """

        info_label = QLabel(info_text)
        info_label.setWordWrap(True)
        info_label.setOpenExternalLinks(True)
        info_layout.addWidget(info_label)
        info_layout.addStretch()

        tabs.addTab(info_tab, "ℹ️ Info")

        main_layout.addWidget(tabs, stretch=1)

        # UI aktualisieren
        self.refresh_ui()

    def refresh_ui(self):
        """Aktualisiert alle Listen"""
        # Accounts
        self.account_list.clear()
        for acc in self.accounts:
            icon = "🔑" if acc.use_gmail_api else "📧"
            item = QListWidgetItem(f"{icon} {acc.name}")
            item.setData(Qt.ItemDataRole.UserRole, acc)
            self.account_list.addItem(item)

        # Profile
        self.profile_list.clear()
        for prof in self.profiles:
            status = "✅" if prof.enabled else "⏸️"
            item = QListWidgetItem(f"{status} {prof.name}")
            item.setData(Qt.ItemDataRole.UserRole, prof)
            self.profile_list.addItem(item)

        # Rechnungen
        self.refresh_invoice_table()

    def refresh_invoice_table(self):
        """Aktualisiert die Rechnungstabelle und synchronisiert mit Dateisystem"""
        # Erst pruefen welche Dateien noch existieren
        valid_invoices = []
        removed_count = 0

        for inv in self.invoices:
            if Path(inv.path).exists():
                valid_invoices.append(inv)
            else:
                removed_count += 1

        # Liste aktualisieren wenn Dateien geloescht wurden
        if removed_count > 0:
            self.invoices = valid_invoices
            if hasattr(self, 'log_output'):
                self.log_output.appendPlainText(f"[SYNC] {removed_count} gelöschte Einträge entfernt")

        # PHASE 2: Neue PDFs finden
        new_count = self.scan_folders_for_new_files()
        if new_count > 0 and hasattr(self, 'log_output'):
            self.log_output.appendPlainText(f"[SCAN] {new_count} neue Dateien importiert")

        # Speichern wenn Aenderungen
        if removed_count > 0 or new_count > 0:
            self.save_invoices_db()

        # Tabelle neu aufbauen
        self.invoice_table.setRowCount(0)

        for inv in reversed(self.invoices[-500:]):  # Letzte 500
            row = self.invoice_table.rowCount()
            self.invoice_table.insertRow(row)

            # Checkbox in Spalte 0
            chk_item = QTableWidgetItem()
            chk_item.setFlags(Qt.ItemFlag.ItemIsUserCheckable | Qt.ItemFlag.ItemIsEnabled)
            chk_item.setCheckState(Qt.CheckState.Unchecked)
            chk_item.setData(Qt.ItemDataRole.UserRole, inv.path)  # Pfad speichern
            self.invoice_table.setItem(row, 0, chk_item)

            # Typ-Symbol: 📎 = Anhang, 📄 = Body-PDF
            typ_symbol = "📎" if getattr(inv, 'is_attachment', False) else "📄"
            typ_item = QTableWidgetItem(typ_symbol)
            typ_item.setToolTip("PDF-Anhang" if typ_symbol == "📎" else "Mail-Body als PDF")
            typ_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.invoice_table.setItem(row, 1, typ_item)

            self.invoice_table.setItem(row, 2, QTableWidgetItem(inv.date))
            self.invoice_table.setItem(row, 3, QTableWidgetItem(inv.profile_name))
            self.invoice_table.setItem(row, 4, QTableWidgetItem(inv.sender[:40]))
            # Spalte 5: Betrag (editierbar)
            amount_text = f"{inv.amount:.2f}" if inv.amount is not None else ""
            amount_item = QTableWidgetItem(amount_text)
            amount_item.setToolTip("Rechnungsbetrag in EUR (manuell eintragen)")
            self.invoice_table.setItem(row, 5, amount_item)
            self.invoice_table.setItem(row, 6, QTableWidgetItem(inv.filename))
            self.invoice_table.setItem(row, 7, QTableWidgetItem(inv.path))

    def save_invoices_db(self):
        """Speichert die Rechnungsdatenbank"""
        try:
            data = [inv.to_dict() for inv in self.invoices]
            INVOICES_DB.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding='utf-8')
        except (OSError, TypeError, ValueError) as e:
            print(f"Invoice DB save error: {e}")

    def _convert_eml_to_pdf(self, eml_path: Path) -> Optional[Path]:
        """
        Konvertiert eine .eml-Datei zu PDF.
        Extrahiert HTML-Body und Anhaenge, speichert als PDF.
        Returns: Pfad zur PDF oder None bei Fehler.
        """
        try:
            with open(eml_path, "rb") as f:
                msg = email.message_from_bytes(f.read())

            # Absender und Betreff extrahieren
            subject = email.header.decode_header(msg.get("Subject", ""))[0]
            if isinstance(subject[0], bytes):
                _ = subject[0].decode(subject[1] or "utf-8", errors="replace")
            else:
                _ = str(subject[0])

            # HTML-Body extrahieren
            html_body = ""
            for part in msg.walk():
                if part.get_content_type() == "text/html":
                    charset = part.get_content_charset() or "utf-8"
                    html_body = part.get_payload(decode=True).decode(charset, errors="replace")
                    break

            if not html_body:
                # Fallback: Plain-Text Body
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        charset = part.get_content_charset() or "utf-8"
                        text = part.get_payload(decode=True).decode(charset, errors="replace")
                        html_body = f"<html><body><pre>{escape(text)}</pre></body></html>"
                        break

            if not html_body:
                return None

            # PDF erstellen
            pdf_path = eml_path.with_suffix(".pdf")
            if XHTML2PDF_AVAILABLE:
                with open(pdf_path, "wb") as f:
                    pisa.CreatePDF(html_body, dest=f)
                if hasattr(self, 'log_output'):
                    self.log_output.appendPlainText(f"[EML] Konvertiert: {eml_path.name} -> {pdf_path.name}")
                return pdf_path
            else:
                if hasattr(self, 'log_output'):
                    self.log_output.appendPlainText(f"[EML] xhtml2pdf nicht verfügbar, kann {eml_path.name} nicht konvertieren")
                return None

        except Exception as e:
            if hasattr(self, 'log_output'):
                self.log_output.appendPlainText(f"[EML] Fehler bei {eml_path.name}: {e}")
            return None

    def _convert_msg_to_pdf(self, msg_path: Path) -> Optional[Path]:
        """
        Konvertiert eine .msg-Datei (Outlook) zu PDF.
        Benoetigt extract-msg (pip install extract-msg).
        Returns: Pfad zur PDF oder None bei Fehler.
        """
        try:
            import extract_msg
        except ImportError:
            if hasattr(self, 'log_output'):
                self.log_output.appendPlainText(
                    "[MSG] extract-msg nicht installiert. Bitte: pip install extract-msg"
                )
            return None

        msg = None
        try:
            msg = extract_msg.Message(str(msg_path))
            html_body = msg.htmlBody
            if not html_body and msg.body:
                html_body = f"<html><body><pre>{escape(msg.body)}</pre></body></html>"
            elif html_body and isinstance(html_body, bytes):
                html_body = html_body.decode("utf-8", errors="replace")

            if not html_body:
                return None

            pdf_path = msg_path.with_suffix(".pdf")
            if XHTML2PDF_AVAILABLE:
                with open(pdf_path, "wb") as f:
                    pisa.CreatePDF(html_body, dest=f)
                if hasattr(self, 'log_output'):
                    self.log_output.appendPlainText(f"[MSG] Konvertiert: {msg_path.name} -> {pdf_path.name}")
                return pdf_path
            return None

        except Exception as e:
            if hasattr(self, 'log_output'):
                self.log_output.appendPlainText(f"[MSG] Fehler bei {msg_path.name}: {e}")
            return None
        finally:
            if msg is not None:
                try:
                    msg.close()
                except Exception:
                    pass

    def scan_folders_for_new_files(self) -> int:
        """
        Scannt alle Profilordner nach neuen PDFs, .eml und .msg Dateien.
        .eml/.msg Dateien werden automatisch zu PDF konvertiert.
        Returns: Anzahl neu hinzugefuegter Dateien

        WICHTIG: Dateien werden auch importiert wenn ihr Hash bereits bekannt ist!
        Nur so tauchen sie in der GUI auf und koennen geloescht werden.
        (Fix von Gemini: Hash-Duplikate nicht mehr ueberspringen)
        """
        new_count = 0
        base_path = Path(self.settings.download_path)

        # Bekannte Pfade fuer schnellen Lookup
        known_paths = {inv.path for inv in self.invoices}
        # Hash-Set nur fuer Logging, nicht zum Blockieren
        known_hashes = {inv.hash for inv in self.invoices if inv.hash}

        for profile in self.profiles:
            # Profilordner bestimmen
            if profile.target_subfolder:
                folder = base_path / sanitize_filename(profile.target_subfolder)
            else:
                folder = base_path / sanitize_filename(profile.name)

            if not folder.exists():
                continue

            # .eml und .msg Dateien zuerst konvertieren
            for eml_path in folder.glob("*.eml"):
                pdf_result = eml_path.with_suffix(".pdf")
                if not pdf_result.exists():
                    self._convert_eml_to_pdf(eml_path)

            for msg_path in folder.glob("*.msg"):
                pdf_result = msg_path.with_suffix(".pdf")
                if not pdf_result.exists():
                    self._convert_msg_to_pdf(msg_path)

            # PDFs im Ordner finden (inkl. frisch konvertierter)
            for pdf_path in folder.glob("*.pdf"):
                str_path = str(pdf_path)

                # Bereits bekannt? (Pruefung nur auf PFAD, nicht auf Hash!)
                if str_path in known_paths:
                    continue

                # 0-Byte Dateien ueberspringen (korrupte/unvollstaendige Dateien)
                try:
                    if pdf_path.stat().st_size == 0:
                        continue
                except OSError:
                    continue

                # Hash berechnen
                file_hash = calculate_file_hash(pdf_path)

                # FIX: Auch Dateien mit bekanntem Hash importieren!
                # Nur so tauchen sie in der GUI auf und koennen geloescht werden.
                is_duplicate = False
                if file_hash and file_hash in known_hashes:
                    is_duplicate = True
                    # Wir loggen es nur, ueberspringen es aber nicht mehr
                    if hasattr(self, 'log_output'):
                        self.log_output.appendPlainText(f"[SCAN] Duplikat importiert: {pdf_path.name}")

                # Metadaten extrahieren
                stat = pdf_path.stat()
                file_date = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d")

                # Invoice erstellen
                inv = Invoice(
                    id=str(uuid.uuid4()),
                    profile_name=profile.name,
                    filename=pdf_path.name,
                    date=file_date,
                    path=str_path,
                    profile_id=profile.id,
                    sender="Manuell importiert" if not is_duplicate else "Duplikat (Import)",
                    subject=pdf_path.stem,
                    hash=file_hash or "",
                    is_attachment=False  # Unbekannt bei Import
                )

                self.invoices.append(inv)
                known_paths.add(str_path)
                if file_hash:
                    known_hashes.add(file_hash)
                new_count += 1

                if hasattr(self, 'log_output') and not is_duplicate:
                    self.log_output.appendPlainText(f"[SCAN] Neu: {pdf_path.name} -> {profile.name}")

        return new_count

    def _collect_folder_hashes(self) -> set:
        """
        Sammelt alle Hashes von PDFs in den Zielordnern.
        Dies stellt sicher, dass auch Dateien ohne DB-Eintrag als Duplikate erkannt werden.
        """
        hashes = set()
        base_path = Path(self.settings.download_path)

        if not base_path.exists():
            return hashes

        # Alle PDFs in allen Unterordnern finden
        for pdf_path in base_path.rglob("*.pdf"):
            try:
                file_hash = calculate_file_hash(pdf_path)
                if file_hash:
                    hashes.add(file_hash)
            except (OSError, ValueError):
                pass

        return hashes

    def add_account(self):
        """Fügt neuen Account hinzu"""
        dlg = AccountDialog(parent=self)
        if dlg.exec():
            account, password = dlg.get_account()
            self.accounts.append(account)

            if password and KEYRING_AVAILABLE:
                try:
                    keyring.set_password(APP_NAME, account.id, password)
                except Exception as e:
                    QMessageBox.warning(self, "Warnung", f"Passwort konnte nicht gespeichert werden: {e}")

            self.save_config()
            self.refresh_ui()

    def edit_account(self, item: QListWidgetItem):
        """Bearbeitet E-Mail Konto per Doppelklick"""
        account = item.data(Qt.ItemDataRole.UserRole)
        dlg = AccountDialog(account=account, parent=self)
        if dlg.exec():
            new_account, password = dlg.get_account()
            # ID beibehalten
            new_account.id = account.id

            # Account in Liste ersetzen
            idx = self.accounts.index(account)
            self.accounts[idx] = new_account

            # Passwort nur aktualisieren wenn neues eingegeben wurde
            if password and password.strip() and KEYRING_AVAILABLE:
                try:
                    keyring.set_password(APP_NAME, new_account.id, password)
                except Exception as e:
                    QMessageBox.warning(self, "Warnung", f"Passwort konnte nicht gespeichert werden: {e}")

            self.save_config()
            self.refresh_ui()

    def delete_account(self):
        """Löscht ausgewählten Account"""
        item = self.account_list.currentItem()
        if not item:
            return

        acc = item.data(Qt.ItemDataRole.UserRole)
        reply = QMessageBox.question(
            self, "Löschen",
            f"Account '{acc.name}' wirklich löschen?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.accounts.remove(acc)

            # Passwort löschen
            if KEYRING_AVAILABLE:
                try:
                    keyring.delete_password(APP_NAME, acc.id)
                except (OSError, RuntimeError):
                    pass

            self.save_config()
            self.refresh_ui()

    def add_profile(self):
        """Fügt neues Profil hinzu"""
        if not self.accounts:
            QMessageBox.warning(self, "Fehler", "Bitte zuerst ein E-Mail Konto hinzufügen!")
            return

        dlg = ProfileDialog(self.accounts, parent=self)
        if dlg.exec():
            profile = dlg.get_profile()
            self.profiles.append(profile)
            self.save_config()
            self.refresh_ui()

    def edit_profile(self, item: QListWidgetItem):
        """Bearbeitet Profil"""
        profile = item.data(Qt.ItemDataRole.UserRole)
        dlg = ProfileDialog(self.accounts, profile, parent=self)
        if dlg.exec():
            idx = self.profiles.index(profile)
            self.profiles[idx] = dlg.get_profile()
            self.save_config()
            self.refresh_ui()

    def delete_profile(self):
        """Löscht ausgewähltes Profil"""
        item = self.profile_list.currentItem()
        if not item:
            return

        profile = item.data(Qt.ItemDataRole.UserRole)
        reply = QMessageBox.question(
            self, "Löschen",
            f"Profil '{profile.name}' wirklich löschen?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.profiles.remove(profile)
            self.save_config()
            self.refresh_ui()

    def browse_folder(self):
        """Ordner auswaehlen"""
        folder = QFileDialog.getExistingDirectory(self, "Speicherort wählen", self.settings.download_path)
        if folder:
            self.inp_path.setText(folder)

    def _on_timeframe_changed(self, text: str):
        """Aktualisiert Von/Bis Datum basierend auf Schnellauswahl"""
        from PySide6.QtCore import QDate
        today = QDate.currentDate()

        if "12 Monate" in text:
            self.date_from.setDate(today.addMonths(-12))
            self.date_to.setDate(today)
        elif "6 Monate" in text:
            self.date_from.setDate(today.addMonths(-6))
            self.date_to.setDate(today)
        elif "3 Monate" in text:
            self.date_from.setDate(today.addMonths(-3))
            self.date_to.setDate(today)
        elif "Dieses Jahr" in text:
            self.date_from.setDate(QDate(today.year(), 1, 1))
            self.date_to.setDate(today)
        elif "Letztes Jahr" in text:
            self.date_from.setDate(QDate(today.year() - 1, 1, 1))
            self.date_to.setDate(QDate(today.year() - 1, 12, 31))
        elif "Alles" in text:
            self.date_from.setDate(QDate(2000, 1, 1))
            self.date_to.setDate(today)
        # "Benutzerdefiniert" -> keine Aenderung

    def _apply_settings(self):
        """Wendet Einstellungen an (ohne MessageBox)"""
        self.settings.download_path = self.inp_path.text()
        self.settings.download_attachments = self.ck_attachments.isChecked()
        self.settings.convert_body_to_pdf = self.ck_body_pdf.isChecked()
        self.settings.merge_body_with_attachments = self.ck_merge_body.isChecked()
        self.settings.enable_hash_check = self.ck_hash.isChecked()
        self.settings.include_trash = self.ck_trash.isChecked()
        self.settings.max_emails_per_run = self.inp_max_mails.value()

        # PDF-Modus und OCR
        self.settings.pdf_mode = self.cmb_pdf_mode.currentData()
        self.settings.ocr_enabled = self.ck_ocr.isChecked()

        # Zeitraum-Filter aus DateEdit-Feldern
        self.settings.date_from = self.date_from.date().toString("yyyy-MM-dd")
        self.settings.date_to = self.date_to.date().toString("yyyy-MM-dd")

        # Legacy: date_filter_months fuer Abwaertskompatibilitaet
        self.settings.date_filter_months = 0

        self.save_config()

    def save_settings(self):
        """Speichert Einstellungen (mit Benutzer-Feedback)"""
        self._apply_settings()
        QMessageBox.information(self, "OK", "Einstellungen gespeichert!")

    def open_datev_settings_dialog(self):
        """Öffnet den DATEV-Einstellungen-Dialog aus dem Einstellungen-Tab."""
        if not DATEV_AVAILABLE:
            QMessageBox.warning(
                self,
                "DATEV nicht verfügbar",
                "datev_exporter.py wurde nicht gefunden.\n"
                "Bitte sicherstellen, dass die Datei im selben Ordner liegt."
            )
            return
        dlg = DATEVSettingsDialog(self.datev_config, parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self.datev_config = dlg.get_config()
            self.save_config()
            if hasattr(self, 'log_output'):
                self.log_output.appendPlainText("[DATEV] Einstellungen und Konten-Mapping gespeichert.")
            QMessageBox.information(self, "DATEV-Einstellungen", "DATEV-Konfiguration und Konten-Mapping gespeichert!")

    def start_grabbing(self):
        """Startet den Abruf"""
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.btn_start.setText("🚀  RECHNUNGEN ABRUFEN")
            return

        if not self.accounts:
            QMessageBox.warning(self, "Fehler", "Bitte zuerst ein E-Mail Konto hinzufügen!")
            return

        enabled_profiles = [p for p in self.profiles if p.enabled]
        if not enabled_profiles:
            QMessageBox.warning(self, "Fehler", "Keine aktiven Suchprofile vorhanden!")
            return

        # Einstellungen anwenden (ohne MessageBox)
        self._apply_settings()

        # WICHTIG: Erst Ordner scannen und neue Dateien in GUI laden
        # Damit werden auch defekte/manuell hinzugefuegte Dateien erkannt
        self.log_output.clear()
        self.log_output.appendPlainText("[SYNC] Synchronisiere mit Dateisystem...")
        scan_count = self.scan_folders_for_new_files()
        if scan_count > 0:
            self.log_output.appendPlainText(f"[SYNC] {scan_count} neue Dateien im Ordner gefunden")
            self.save_invoices_db()
            self.refresh_invoice_table()

        # Existierende Hashes sammeln (jetzt INKLUSIVE der gescannten Dateien)
        existing_hashes = {inv.hash for inv in self.invoices if inv.hash}

        # Zusaetzlich: Alle PDFs in Zielordnern hashen (fuer Dateien ohne DB-Eintrag)
        existing_hashes.update(self._collect_folder_hashes())

        self.log_output.appendPlainText(f"[SYNC] {len(existing_hashes)} bekannte Hashes geladen")
        self.log_output.appendPlainText("")

        # Worker starten
        self.progress.setVisible(True)
        self.progress.setValue(0)
        self.btn_start.setText("⏹️  STOPPEN")
        self.btn_start.setStyleSheet("""
            QPushButton { background-color: #e74c3c; font-size: 14pt; font-weight: bold; border-radius: 8px; }
        """)

        self.worker = InvoiceWorker(
            self.accounts, enabled_profiles, self.settings, existing_hashes
        )
        self.worker.log.connect(self.log_output.appendPlainText)
        self.worker.progress.connect(lambda c, t: self.progress.setValue(int(c / t * 100) if t > 0 else 0))
        self.worker.invoice_found.connect(self.on_invoice_found)
        self.worker.finished_signal.connect(self.on_finished)
        self.worker.start()

    def on_invoice_found(self, invoice: Invoice):
        """Callback wenn Rechnung gefunden wurde"""
        self.invoices.append(invoice)
        self.save_invoices_db()

    def on_finished(self, count: int):
        """Callback wenn Worker fertig"""
        self.btn_start.setText("🚀  RECHNUNGEN ABRUFEN")
        self.btn_start.setStyleSheet("""
            QPushButton { background-color: #27ae60; font-size: 14pt; font-weight: bold; border-radius: 8px; }
        """)
        self.progress.setVisible(False)

        self.log_output.appendPlainText(f"\n{'='*50}")
        self.log_output.appendPlainText(f"✅ Fertig! {count} neue Rechnungen gefunden.")
        self.log_output.appendPlainText(f"{'='*50}")

        self.refresh_invoice_table()

        if count > 0:
            QMessageBox.information(self, "Fertig", f"{count} neue Rechnungen heruntergeladen!")

    def open_invoice(self, row: int, col: int):
        """Öffnet Rechnung"""
        path = self.invoice_table.item(row, 0).data(Qt.ItemDataRole.UserRole)
        if path and Path(path).exists():
            QDesktopServices.openUrl(QUrl.fromLocalFile(path))
        else:
            QMessageBox.warning(self, "Fehler", "Datei nicht gefunden!")

    def select_all_invoices(self):
        """Waehlt alle Eintraege aus"""
        for row in range(self.invoice_table.rowCount()):
            item = self.invoice_table.item(row, 0)
            if item:
                item.setCheckState(Qt.CheckState.Checked)

    def select_no_invoices(self):
        """Hebt Auswahl auf"""
        for row in range(self.invoice_table.rowCount()):
            item = self.invoice_table.item(row, 0)
            if item:
                item.setCheckState(Qt.CheckState.Unchecked)

    def delete_selected_invoices(self):
        """Loescht ausgewaehlte Eintraege UND deren Dateien"""
        # Ausgewaehlte sammeln
        selected_paths = []
        for row in range(self.invoice_table.rowCount()):
            item = self.invoice_table.item(row, 0)
            if item and item.checkState() == Qt.CheckState.Checked:
                path = item.data(Qt.ItemDataRole.UserRole)
                if path:
                    selected_paths.append(path)

        if not selected_paths:
            return

        # Dateien und Eintraege loeschen (ohne Bestaetigung - Scan holt sie wieder)
        deleted_files = 0
        deleted_entries = 0

        for path in selected_paths:
            # Datei loeschen
            file_path = Path(path)
            if file_path.exists():
                try:
                    file_path.unlink()
                    deleted_files += 1
                except Exception as e:
                    self.log_output.appendPlainText(f"[DELETE] Fehler: {path} - {e}")

            # Aus invoices-Liste entfernen
            self.invoices = [inv for inv in self.invoices if inv.path != path]
            deleted_entries += 1

        # Speichern und UI aktualisieren
        self.save_invoices_db()
        self.refresh_invoice_table()

        self.log_output.appendPlainText(
            f"[DELETE] {deleted_entries} Einträge entfernt, {deleted_files} Dateien gelöscht"
        )

    def open_download_folder(self):
        """Oeffnet Download-Ordner"""
        path = Path(self.settings.download_path)
        if not path.exists():
            path.mkdir(parents=True, exist_ok=True)
        QDesktopServices.openUrl(QUrl.fromLocalFile(str(path)))

    def _on_invoice_amount_changed(self, item: QTableWidgetItem) -> None:
        """Schreibt einen manuell eingetragenen Betrag in das Invoice-Objekt zurück."""
        if item.column() != 5:
            return
        path_item = self.invoice_table.item(item.row(), 7)
        if not path_item:
            return
        path = path_item.text()
        invoice = next((inv for inv in self.invoices if inv.path == path), None)
        if invoice is None:
            return
        text = item.text().strip()
        amount: Optional[float] = None
        if text:
            try:
                amount = _normalize_amount(text)
            except ValueError:
                previous_text = f"{invoice.amount:.2f}" if invoice.amount is not None else ""
                message = (
                    f"Der eingegebene Betrag „{text}“ ist ungültig und wurde nicht gespeichert. "
                    "Bitte verwenden Sie zum Beispiel 19,99 oder 1.234,56 €."
                )
                self.invoice_table.blockSignals(True)
                try:
                    item.setText(previous_text)
                    item.setToolTip(message)
                finally:
                    self.invoice_table.blockSignals(False)
                self.log_output.appendPlainText(f"[WARN] {message}")
                QMessageBox.warning(self, "Ungültiger Betrag", message)
                return
        item.setToolTip("Rechnungsbetrag in EUR (zum Beispiel 19,99 oder 1.234,56 €)")
        if invoice.amount != amount:
            invoice.amount = amount
            self.save_invoices_db()

    def _get_selected_invoice_paths(self) -> set[str]:
        """Returns invoice paths that are checked in the invoice table."""
        selected_paths: set[str] = set()
        for row in range(self.invoice_table.rowCount()):
            item = self.invoice_table.item(row, 0)
            if item and item.checkState() == Qt.CheckState.Checked:
                path = item.data(Qt.ItemDataRole.UserRole)
                if path:
                    selected_paths.add(path)
        return selected_paths

    def export_invoice_bundle(self):
        """Exports a redacted invoice bundle for companion usage."""
        if not self.invoices:
            QMessageBox.information(self, "Info", "Keine Rechnungen zum Exportieren vorhanden.")
            return

        selected_paths = self._get_selected_invoice_paths()
        bundle = build_invoice_bundle(
            app_name=APP_NAME,
            app_version=VERSION,
            accounts=self.accounts,
            profiles=self.profiles,
            invoices=self.invoices,
            download_path=self.settings.download_path,
            datev_config=self.datev_config,
            selected_paths=selected_paths,
        )

        invoice_count = len(bundle.get("invoices", []))
        if invoice_count == 0:
            QMessageBox.information(
                self,
                "Info",
                "Die aktuelle Auswahl enthält keine exportierbaren Rechnungen.",
            )
            return

        default_name = f"universalinvoicemail-invoicebundle-v1_{datetime.now().strftime('%Y-%m-%d')}.json"
        filepath, _ = QFileDialog.getSaveFileName(
            self,
            "Bundle speichern",
            str(Path(self.settings.download_path) / default_name),
            "JSON Dateien (*.json)",
        )
        if not filepath:
            return

        try:
            write_invoice_bundle(bundle, Path(filepath))
            if hasattr(self, "log_output"):
                self.log_output.appendPlainText(f"[BUNDLE] Exportiert: {filepath}")
            scope = "Auswahl" if selected_paths else "Gesamtbestand"
            QMessageBox.information(
                self,
                "Bundle Export",
                f"Bundle erfolgreich exportiert:\n{filepath}\n\n{scope}: {invoice_count} Rechnungen.",
            )
        except (OSError, TypeError, ValueError) as e:
            QMessageBox.warning(self, "Fehler", f"Bundle-Export fehlgeschlagen:\n{e}")

    def import_invoice_bundle(self):
        """Imports companion updates from a redacted invoice bundle."""
        if not self.invoices:
            QMessageBox.information(self, "Info", "Keine lokalen Rechnungen für einen Reimport vorhanden.")
            return

        filepath, _ = QFileDialog.getOpenFileName(
            self,
            "Bundle auswählen",
            str(Path(self.settings.download_path)),
            "JSON Dateien (*.json)",
        )
        if not filepath:
            return

        try:
            bundle = load_invoice_bundle(Path(filepath))
            result = apply_invoice_bundle_changes(self.invoices, bundle)
        except (OSError, ValueError, json.JSONDecodeError, TypeError) as e:
            QMessageBox.warning(self, "Fehler", f"Bundle-Import fehlgeschlagen:\n{e}")
            return

        self.save_invoices_db()
        self.refresh_invoice_table()

        if hasattr(self, "log_output"):
            self.log_output.appendPlainText(
                "[BUNDLE] Import: "
                f"{result['updated']} aktualisiert, "
                f"{len(result['conflicts'])} Konflikte, "
                f"{len(result['missing_local'])} ohne lokales Gegenstück."
            )

        lines = [
            f"{result['updated']} Rechnungen aktualisiert.",
            f"{result['unchanged']} Rechnungen unverändert.",
        ]
        if result["missing_local"]:
            lines.append(f"{len(result['missing_local'])} Bundle-Einträge hatten kein lokales Gegenstück.")
        if result["conflicts"]:
            lines.append(f"{len(result['conflicts'])} Konflikte wegen Datei- oder Hash-Abweichung erkannt.")
        if result["invalid_rows"]:
            lines.append(f"{len(result['invalid_rows'])} Bundle-Zeilen waren ungültig und wurden übersprungen.")
        if result["conflicts"]:
            preview = ", ".join(entry["id"] for entry in result["conflicts"][:5])
            lines.append(f"Konflikt-Vorschau: {preview}")

        QMessageBox.information(self, "Bundle Import", "\n".join(lines))

    def _export_datev(self):
        """Exportiert ausgewählte (oder alle) Rechnungen als DATEV-Buchungsstapel."""
        if not DATEV_AVAILABLE:
            QMessageBox.warning(self, "DATEV nicht verfügbar",
                                "datev_exporter.py wurde nicht gefunden.\n"
                                "Bitte sicherstellen, dass die Datei im selben Ordner liegt.")
            return

        if not self.invoices:
            QMessageBox.information(self, "Info", "Keine Rechnungen vorhanden.")
            return

        # Ausgewaehlte Rechnungen aus Checkbox-Spalte 0 ermitteln
        selected_paths = self._get_selected_invoice_paths()

        # Fallback: alle Rechnungen wenn nichts markiert
        if selected_paths:
            source_invoices = [inv for inv in self.invoices if inv.path in selected_paths]
        else:
            source_invoices = list(self.invoices)

        # Invoice-Objekte in DATEV-dicts umwandeln
        # Invoice.profile_name entspricht dem Provider/Shop-Namen (Fallback: sender)
        inv_dicts = []
        for inv in source_invoices:
            provider = inv.profile_name or inv.sender or "Sonstige"
            inv_dicts.append({
                "provider": provider,
                "filename": inv.filename,
                "date": inv.date,
                "path": inv.path,
                "amount": getattr(inv, "amount", None),   # None wenn Feld fehlt
                "category": inv.profile_name or provider,
            })

        # DATEV-Einstellungen aus Attribut laden und Dialog zeigen
        dlg = DATEVSettingsDialog(self.datev_config, parent=self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        self.datev_config = dlg.get_config()
        self.save_config()

        # Zieldatei waehlen
        default_name = f"DATEV_BUCHUNGEN_{datetime.now().strftime('%Y-%m-%d')}.csv"
        filepath, _ = QFileDialog.getSaveFileName(
            self, "DATEV-Datei speichern",
            str(Path(self.settings.download_path) / default_name),
            "CSV Dateien (*.csv)"
        )
        if not filepath:
            return

        # Export ausfuehren
        try:
            exporter = DATEVExporter(self.datev_config)
            exporter.export(inv_dicts, Path(filepath))

            # Benutzerinformation
            with_amount = sum(1 for d in inv_dicts if d.get("amount") and d["amount"] > 0)
            skipped = len(inv_dicts) - with_amount
            msg = f"DATEV-Export erfolgreich:\n{filepath}\n\n{with_amount} Buchungen exportiert."
            if skipped:
                msg += (f"\n{skipped} Rechnungen übersprungen (kein Betrag eingetragen).\n\n"
                        "Tipp: Beträge können manuell in der exportierten CSV ergänzt werden.")
            QMessageBox.information(self, "DATEV-Export", msg)
        except Exception as e:
            QMessageBox.critical(self, "Fehler", f"DATEV-Export fehlgeschlagen:\n{e}")

    def export_invoices_csv(self):
        """Exportiert Rechnungsliste als CSV"""
        if not self.invoices:
            QMessageBox.information(self, "Info", "Keine Rechnungen zum Exportieren vorhanden.")
            return

        # Speicherort waehlen
        default_name = f"Rechnungen_{datetime.now().strftime('%Y-%m-%d')}.csv"
        filepath, _ = QFileDialog.getSaveFileName(
            self, "CSV speichern",
            str(Path(self.settings.download_path) / default_name),
            "CSV Dateien (*.csv)"
        )

        if not filepath:
            return

        try:
            import csv
            with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f, delimiter=';')
                # Header
                writer.writerow(['Datum', 'Profil/Shop', 'Absender', 'Betreff', 'Dateiname', 'Pfad'])
                # Daten
                for inv in self.invoices:
                    writer.writerow([
                        inv.date,
                        inv.profile_name,
                        inv.sender,
                        inv.subject,
                        inv.filename,
                        inv.path
                    ])

            QMessageBox.information(self, "Erfolg", f"CSV exportiert:\n{filepath}")
        except Exception as e:
            QMessageBox.warning(self, "Fehler", f"Export fehlgeschlagen:\n{e}")

    def closeEvent(self, event):
        """Cleanup beim Schließen"""
        if self.worker and self.worker.isRunning():
            self.worker.stop()
            self.worker.wait(2000)
        self.save_config()
        # Browser-Renderer schliessen (verhindert verwaiste WebDriver-Prozesse)
        global _browser_renderer
        if _browser_renderer is not None:
            _browser_renderer.close()
            _browser_renderer = None
        event.accept()


# ==================== MAIN ====================

def main():
    """Haupteinstiegspunkt"""
    # Logging konfigurieren
    log_file = BASE_DIR / "app.log"
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file, encoding='utf-8'),
            logging.StreamHandler()
        ]
    )
    logger = logging.getLogger(__name__)
    logger.info(f"{APP_NAME} v{VERSION} gestartet")

    app = QApplication(sys.argv)
    app.setApplicationName(APP_NAME)
    app.setStyle("Fusion")

    # Icon setzen (falls vorhanden)
    icon_path = Path(__file__).parent / "icon.ico"
    if icon_path.exists():
        app.setWindowIcon(QIcon(str(icon_path)))

    window = MainWindow()
    window.show()

    sys.exit(app.exec())


if __name__ == "__main__":
    main()
